import os
import re
import sys
import datetime
import logging
from decimal import Decimal, InvalidOperation
from tkinter import filedialog, messagebox
import tempfile 

# PDF Imports
import pdfplumber
from reportlab.lib.styles import ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer, Image # ReportLab Table
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter # letter is not folio, ensure folio is used or defined
from reportlab.lib.units import inch
import camelot
import fitz # PyMuPDF
import random

# Word Imports
from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.table import Table as DocxTable # Alias for docx.table.Table to avoid conflict

class FileHandler:
    def __init__(self, variables, title_var, date_var, logo_path_var, address_var, prepared_by_var, noted_by_var_1, noted_by_var_2, checked_by_var):
        self.variables = variables
        self.title_var = title_var
        self.date_var = date_var
        self.logo_path_var = logo_path_var
        self.address_var = address_var
        self.prepared_by_var = prepared_by_var
        self.noted_by_var_1 = noted_by_var_1
        self.noted_by_var_2 = noted_by_var_2
        self.checked_by_var = checked_by_var
        # Ensure calculator is accessible if needed for recalculation after load
        self.calculator = variables.get('calculator') if 'calculator' in variables else None


    def parse_amount(self, text):
        if not text or str(text).strip() == "":
            return ""
        try:
            # Remove currency symbols, spaces, and all but one decimal point
            cleaned_text = re.sub(r'[^\d.]', '', str(text).replace(',', ''))
            if not cleaned_text: # Handle cases like "₱" becoming ""
                return ""
            return str(Decimal(cleaned_text))
        except InvalidOperation: # Catch specific Decimal error for non-numeric after cleaning
            logging.warning(f"Could not parse amount '{text}' due to InvalidOperation after cleaning to '{cleaned_text}'")
            return "0.00" # Or handle as error, or return original text
        except Exception as e:
            logging.warning(f"Could not parse amount '{text}': {e}")
            return str(text) # Return original text if parsing fails

    def format_date_for_display(self, date_str):
        try:
            date_obj = datetime.datetime.strptime(date_str, "%m/%d/%Y")
            return date_obj.strftime("%B %d, %Y")
        except ValueError:
            logging.warning(f"Invalid date format for display: {date_str}")
            return date_str

    # --- Internal PDF Content Generation ---
    def _build_pdf_elements(self):
        elements = []
        
        def format_amount_pdf(value_var_or_str):
            val_str = ""
            if hasattr(value_var_or_str, 'get'): # Check if it's a StringVar
                val_str = value_var_or_str.get()
            else:
                val_str = str(value_var_or_str)

            if val_str:
                try:
                    cleaned_val_str = str(val_str).replace(',', '')
                    if not cleaned_val_str: return ""
                    amount = Decimal(cleaned_val_str)
                    return f"{amount:,.2f}"
                except Exception as e:
                    logging.warning(f"Could not format amount '{val_str}' for PDF: {e}")
                    return str(val_str)
            return ""

        styles = getSampleStyleSheet()
        folio_size = (8.5 * inch, 13 * inch) 
        doc_left_margin = 0.5 * inch
        doc_right_margin = 0.5 * inch

        addressStyle = ParagraphStyle(name='addressStyle', fontName='Helvetica', fontSize=10, leading=12, alignment=1) # Center
        titleStyle = ParagraphStyle(name='titleStyle', fontName='Helvetica-Bold', fontSize=12, leading=14, alignment=1) # Center
        dateStyle = ParagraphStyle(name='dateStyle', fontName='Helvetica', fontSize=8, leading=10, alignment=1, spaceBefore=4) # Center
        tableboldStyle = ParagraphStyle(name='tableBoldStyle', fontName='Helvetica-Bold', fontSize=10, leading=12, spaceAfter=4, spaceBefore=6) 
        footerStyle = ParagraphStyle(name='footerStyle', fontName='Helvetica', fontSize=8, leading=10, alignment=1) # Center
        notedStyle = ParagraphStyle(name='notedStyle', fontName='Helvetica', fontSize=8, leading=10, alignment=1) # Center

        logo_path = self.logo_path_var.get()
        address_text = self.address_var.get() or " "

        logo_img = None
        logo_placeholder_text = ""
        if logo_path and os.path.exists(logo_path):
            try:
                logo_img = Image(logo_path, width=1.18 * inch, height=1.18 * inch) # Fixed size from DOCX
                logo_img.hAlign = 'CENTER'
            except Exception as e:
                logging.warning(f"Could not load or process logo image '{logo_path}': {e}")
                logo_placeholder_text = "[Logo Error]"
        elif logo_path:
            logo_placeholder_text = "[Logo N/A]"
        
        logo_cell_content = Paragraph(logo_placeholder_text, styles['Italic']) if logo_img is None else logo_img

        header_text_elements = [
            Paragraph(address_text, addressStyle),
            Spacer(1, 4), 
            Paragraph("CASH FLOW STATEMENT", titleStyle),
            Paragraph(f"For the Month of {self.format_date_for_display(self.date_var.get())}", dateStyle)
        ]
        
        page_width = folio_size[0] - doc_left_margin - doc_right_margin
        logo_col_width = 1.58 * inch 
        text_col_width = page_width - logo_col_width - (0.05*inch) 

        header_table_data = [[logo_cell_content, header_text_elements]]
        header_table = Table(header_table_data, colWidths=[logo_col_width, text_col_width], hAlign='LEFT') # ReportLab Table
        header_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (0,0), 'MIDDLE'), 
            ('ALIGN', (0,0), (0,0), 'CENTER'),   
            ('VALIGN', (1,0), (1,0), 'TOP'),    
            ('ALIGN', (1, 0), (1, 0), 'CENTER'), 
            ('LEFTPADDING', (0,0), (-1,-1), 0), ('RIGHTPADDING', (0,0), (-1,-1), 0),
            ('BOTTOMPADDING', (0,0), (-1,-1), 0), ('TOPPADDING', (0,0), (-1,-1), 0),
        ]))
        elements.append(header_table)
        elements.append(Spacer(1, 8)) 

        data_label_width = page_width * 0.70 
        data_value_width = page_width * 0.30 
        
        common_table_style_list = [
            ('BOX', (0,0), (-1,-1), 0.5, colors.black), 
            ('INNERGRID', (0,0), (-1,-1), 0.5, colors.black), 
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('ALIGN', (1, 0), (1, -1), 'RIGHT'), 
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0,0), (0,-1), 5), 
            ('RIGHTPADDING', (1,0), (1,-1), 5), 
            ('TOPPADDING', (0,0), (-1,-1), 2), # Reduced padding
            ('BOTTOMPADDING', (0,0), (-1,-1), 2), # Reduced padding
        ]
        
        # Beginning Cash Balances
        elements.append(Paragraph("Beginning Cash Balances", tableboldStyle))
        beg_data = [
            ["Cash in Bank-beg", format_amount_pdf(self.variables['cash_bank_beg'])],
            ["Cash on Hand-beg", format_amount_pdf(self.variables['cash_hand_beg'])]
        ]
        beg_table = Table(beg_data, colWidths=[data_label_width, data_value_width]) # ReportLab Table
        beg_table.setStyle(TableStyle(common_table_style_list))
        elements.append(beg_table)
        elements.append(Spacer(1, 4))

        # Cash Inflows
        elements.append(Paragraph("Cash Inflows", tableboldStyle))
        inflows_data = [
            ["Monthly Dues Collected", format_amount_pdf(self.variables['monthly_dues'])],
            ["Certifications Issued", format_amount_pdf(self.variables['certifications'])],
            ["Membership Fee", format_amount_pdf(self.variables['membership_fee'])],
            ["Vehicle Stickers", format_amount_pdf(self.variables['vehicle_stickers'])],
            ["Rentals", format_amount_pdf(self.variables['rentals'])],
            ["Solicitations/Donations", format_amount_pdf(self.variables['solicitations'])],
            ["Interest Income on Bank Deposits", format_amount_pdf(self.variables['interest_income'])],
            ["Livelihood Management Fee", format_amount_pdf(self.variables['livelihood_fee'])],
            ["Withdrawal from Bank", format_amount_pdf(self.variables['withdrawal_from_bank'])], 
            ["Others (Inflow)", format_amount_pdf(self.variables['inflows_others'])],
        ]
        total_receipts_row = ["Total Cash Receipts", format_amount_pdf(self.variables['total_receipts'])]
        inflows_table_data = inflows_data + [total_receipts_row]
        
        inflows_table = Table(inflows_table_data, colWidths=[data_label_width, data_value_width]) # ReportLab Table
        inflows_table_style = TableStyle(common_table_style_list + [('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold')]) 
        inflows_table.setStyle(inflows_table_style)
        elements.append(inflows_table)
        elements.append(Spacer(1, 4))

        # Cash Outflows
        elements.append(Paragraph("Less: Cash Outflows", tableboldStyle))
        outflows_data = [
            ["Snacks/Meals for Visitors", format_amount_pdf(self.variables['snacks_meals'])],
            ["Transportation Expenses", format_amount_pdf(self.variables['transportation'])],
            ["Office Supplies Expense", format_amount_pdf(self.variables['office_supplies'])],
            ["Printing and Photocopy", format_amount_pdf(self.variables['printing'])],
            ["Labor", format_amount_pdf(self.variables['labor'])],
            ["Billboard Expense", format_amount_pdf(self.variables['billboard'])],
            ["Clearing/Cleaning Charges", format_amount_pdf(self.variables['cleaning'])],
            ["Miscellaneous Expenses", format_amount_pdf(self.variables['misc_expenses'])],
            ["Federation Fee", format_amount_pdf(self.variables['federation_fee'])],
            ["HOA-BOD Uniforms", format_amount_pdf(self.variables['uniforms'])],
            ["BOD Meeting", format_amount_pdf(self.variables['bod_mtg'])],
            ["General Assembly", format_amount_pdf(self.variables['general_assembly'])],
            ["Cash Deposit to Bank", format_amount_pdf(self.variables['cash_deposit'])],
            ["Withholding Tax on Bank Deposit", format_amount_pdf(self.variables['withholding_tax'])],
            ["Refund", format_amount_pdf(self.variables['refund'])],
            ["Others (Outflow)", format_amount_pdf(self.variables['outflows_others'])],
        ]
        total_outflows_row = ["Total Cash Outflows/Disbursements", format_amount_pdf(self.variables['cash_outflows'])]
        outflows_table_data = outflows_data + [total_outflows_row]

        outflows_table = Table(outflows_table_data, colWidths=[data_label_width, data_value_width]) # ReportLab Table
        outflows_table_style = TableStyle(common_table_style_list + [('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold')]) 
        outflows_table.setStyle(outflows_table_style)
        elements.append(outflows_table)
        elements.append(Spacer(1, 4))
        
        # Ending Cash Balance
        elements.append(Paragraph("Ending Cash Balance", tableboldStyle))
        ending_data = [["Ending Cash Balance", format_amount_pdf(self.variables['ending_cash'])]]
        ending_table = Table(ending_data, colWidths=[data_label_width, data_value_width]) # ReportLab Table
        ending_table_style = TableStyle(common_table_style_list + [('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold')]) 
        ending_table.setStyle(ending_table_style)
        elements.append(ending_table)
        elements.append(Spacer(1, 4))

        # Breakdown of Cash
        elements.append(Paragraph("Breakdown of Cash", tableboldStyle))
        breakdown_data = [
            ["Cash in Bank", format_amount_pdf(self.variables['ending_cash_bank'])],
            ["Cash on Hand", format_amount_pdf(self.variables['ending_cash_hand'])]
        ]
        breakdown_table = Table(breakdown_data, colWidths=[data_label_width, data_value_width]) # ReportLab Table
        breakdown_table.setStyle(TableStyle(common_table_style_list))
        elements.append(breakdown_table)
        elements.append(Spacer(1, 12)) 
        
        # Footer Signatories
        prepared_name = self.prepared_by_var.get() or "_______________________"
        noted_name_1 = self.noted_by_var_1.get() or "_______________________"
        noted_name_2 = self.noted_by_var_2.get() or "_______________________"
        checked_name = self.checked_by_var.get() or "_______________________"
        
        sign_col_width = (page_width / 2) - (0.05 * inch) 
        
        # Prepared by / Checked by
        prep_check_data = [
            [Paragraph(f"Prepared by:<br/><br/><b>{prepared_name}</b><br/>HOA Treasurer", footerStyle)],
            [Paragraph(f"Checked by:<br/><br/><b>{checked_name}</b><br/>HOA Auditor", footerStyle)]
        ]
        prep_check_table = Table(prep_check_data, colWidths=[sign_col_width, sign_col_width]) # ReportLab Table
        prep_check_table.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'CENTER'), ('VALIGN', (0,0),(-1,-1), 'TOP'),
            ('LEFTPADDING', (0,0),(-1,-1),0), ('RIGHTPADDING', (0,0),(-1,-1),0),
            ('TOPPADDING', (0,0), (-1,-1), 0), ('BOTTOMPADDING', (0,0), (-1,-1), 0)
        ]))
        elements.append(prep_check_table)
        elements.append(Spacer(1, 12)) 
        
        elements.append(Paragraph("Noted by:", notedStyle)) 
        elements.append(Spacer(1, 4)) 
        
        noted_data = [
            [Paragraph(f"<b>{noted_name_1}</b><br/>HOA President", footerStyle)],
            [Paragraph(f"<b>{noted_name_2}</b><br/>CHUDD HCD-CORDS", footerStyle)]
        ]
        noted_table = Table(noted_data, colWidths=[sign_col_width, sign_col_width]) # ReportLab Table
        noted_table.setStyle(TableStyle([
            ('ALIGN', (0,0), (-1,-1), 'CENTER'), ('VALIGN', (0,0),(-1,-1), 'TOP'),
            ('LEFTPADDING', (0,0),(-1,-1),0), ('RIGHTPADDING', (0,0),(-1,-1),0),
            ('TOPPADDING', (0,0), (-1,-1), 0), ('BOTTOMPADDING', (0,0), (-1,-1), 0)
        ]))
        elements.append(noted_table)
        
        return elements

    def _create_pdf_at_path(self, filename):
        try:
            folio_size = (8.5 * inch, 13 * inch)
            doc = SimpleDocTemplate(
                filename,
                pagesize=folio_size,
                topMargin=0.5*inch,
                bottomMargin=0.5*inch,
                leftMargin=0.5*inch,
                rightMargin=0.5*inch
            )
            elements = self._build_pdf_elements()
            doc.build(elements)
            return True
        except Exception as e:
            logging.exception(f"Error creating PDF content for {filename}")
            return False

    def export_to_pdf(self):
        try:
            default_filename = f"cash_flow_statement_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            filename = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf")],
                initialfile=default_filename,
                title="Save PDF As"
            )
            if not filename:
                return {"status": "cancelled", "message": "Export to PDF cancelled by user."}

            if self._create_pdf_at_path(filename):
                logging.info(f"PDF successfully exported to {filename}")
                return {"status": "success", "message": f"PDF successfully exported to {filename}", "filename": filename}
            else:
                return {"status": "error", "message": f"Failed to create PDF at {filename}.\nCheck logs for details and ensure ReportLab is installed."}
        except Exception as e:
            logging.exception("Error during PDF export process")
            return {"status": "error", "message": f"An unexpected error occurred during PDF export: {str(e)}"}

    def generate_temp_pdf(self):
        try:
            fd, temp_filename = tempfile.mkstemp(suffix=".pdf", prefix="cash_flow_")
            os.close(fd) 
            if self._create_pdf_at_path(temp_filename):
                logging.info(f"Temporary PDF for email generated at {temp_filename}")
                return {"status": "success", "filename": temp_filename}
            else:
                if os.path.exists(temp_filename):
                    try: os.remove(temp_filename)
                    except Exception as e_del: logging.error(f"Could not remove temp PDF {temp_filename} after generation failure: {e_del}")
                return {"status": "error", "message": "Failed to generate temporary PDF."}
        except Exception as e:
            logging.exception("Error generating temporary PDF")
            return {"status": "error", "message": f"An unexpected error occurred while generating temporary PDF: {str(e)}"}

    def _build_docx_document(self):
        doc = Document()
        
        def format_amount_docx(value_var_or_str):
            val_str = ""
            if hasattr(value_var_or_str, 'get'): 
                val_str = value_var_or_str.get()
            else:
                val_str = str(value_var_or_str)

            if val_str:
                try:
                    cleaned_val_str = str(val_str).replace(',', '')
                    if not cleaned_val_str: return ""
                    amount = Decimal(cleaned_val_str)
                    return f"{amount:,.2f}"
                except Exception as e:
                    logging.warning(f"Could not format amount '{val_str}' for Word: {e}")
                    return str(val_str) 
            return ""
        
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(13)
        section.top_margin = Inches(0.5); section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.5); section.right_margin = Inches(0.5)

        header = section.header
        header.is_linked_to_previous = False
        
        hdr_elm = header._element 
        for child_elm in list(hdr_elm): 
            hdr_elm.remove(child_elm)
        
        page_content_width = section.page_width - section.left_margin - section.right_margin

        logo_path = self.logo_path_var.get()
        address_text = self.address_var.get() or " "

        header_table = header.add_table(rows=1, cols=2, width=page_content_width) 
        header_table.autofit = False; header_table.allow_autofit = False
        
        logo_col_width_val = Inches(1.58)
        text_col_width_val = page_content_width - logo_col_width_val
        if text_col_width_val < Inches(0.5): text_col_width_val = Inches(0.5)
        
        header_table.columns[0].width = logo_col_width_val
        header_table.columns[1].width = text_col_width_val
        
        logo_cell = header_table.cell(0, 0)
        text_cell = header_table.cell(0, 1)
        logo_cell.width = logo_col_width_val; text_cell.width = text_col_width_val
        logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        text_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP 

        tc_logo_elm = logo_cell._tc
        for child_elm in list(tc_logo_elm):
            tc_logo_elm.remove(child_elm)
        
        logo_para = logo_cell.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if logo_path and os.path.exists(logo_path):
            try:
                # Separate run creation for clarity in case of add_picture failure
                run = logo_para.add_run()
                run.add_picture(logo_path, width=Inches(1.18), height=Inches(1.18)) 
            except Exception as e:
                logging.warning(f"Could not add picture {logo_path} to DOCX header: {e}")
                # Clean the paragraph and add error text
                p_elm = logo_para._p
                for r_elm in list(p_elm.xpath('./w:r')): # Remove any existing runs
                    p_elm.remove(r_elm)
                logo_para.add_run("[Logo Error]").italic = True
        elif logo_path: # logo_path is not empty but file doesn't exist
            logo_para.add_run("[Logo N/A]").italic = True
        # If logo_path is empty, logo_para remains an empty, centered paragraph.

        tc_text_elm = text_cell._tc
        for child_elm in list(tc_text_elm):
            tc_text_elm.remove(child_elm)

        p_address = text_cell.add_paragraph()
        run_address = p_address.add_run(address_text); run_address.font.name = 'Helvetica'; run_address.font.size = Pt(10)
        p_address.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_address.paragraph_format.space_after = Pt(4)
        
        p_title = text_cell.add_paragraph()
        run_title = p_title.add_run("CASH FLOW STATEMENT"); run_title.font.name = 'Helvetica'; run_title.bold = True; run_title.font.size = Pt(12)
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_title.paragraph_format.space_after = Pt(0)

        p_date = text_cell.add_paragraph()
        run_date = p_date.add_run(f"For the Month of {self.format_date_for_display(self.date_var.get())}"); run_date.font.name = 'Helvetica'; run_date.font.size = Pt(8)
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_date.paragraph_format.space_before = Pt(0); p_date.paragraph_format.space_after = Pt(4)

        def set_cell_style(cell, text, size=8, bold=False, align='left', font='Helvetica'):
            tc_element = cell._tc 
            for p_element in tc_element.xpath('./w:p'): 
                tc_element.remove(p_element)
            
            para = cell.add_paragraph() 
            run = para.add_run(text)    
            
            run.font.name = font; run.font.size = Pt(size); run.bold = bold
            if align == 'right': para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == 'center': para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else: para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(1); para.paragraph_format.space_after = Pt(1) 
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        table_content_width = section.page_width - section.left_margin - section.right_margin
        # Ensure calculated widths are integers (EMUs)
        table_col1_width = int(table_content_width * 0.70)
        table_col2_width = int(table_content_width * 0.30)

        title_para_fmt = {"space_before": Pt(4), "space_after": Pt(2)}

        p_beg_title = doc.add_paragraph(); run_beg_title = p_beg_title.add_run("Beginning Cash Balances")
        run_beg_title.font.name = 'Helvetica'; run_beg_title.bold = True; run_beg_title.font.size = Pt(10)
        p_beg_title.paragraph_format.space_before = title_para_fmt["space_before"]; p_beg_title.paragraph_format.space_after = title_para_fmt["space_after"]
        beg_table = doc.add_table(rows=2, cols=2); beg_table.style = 'Table Grid'; beg_table.autofit = False
        beg_table.columns[0].width = table_col1_width; beg_table.columns[1].width = table_col2_width
        set_cell_style(beg_table.cell(0,0), "Cash in Bank-beg"); set_cell_style(beg_table.cell(0,1), format_amount_docx(self.variables['cash_bank_beg']), align='right')
        set_cell_style(beg_table.cell(1,0), "Cash on Hand-beg"); set_cell_style(beg_table.cell(1,1), format_amount_docx(self.variables['cash_hand_beg']), align='right')

        p_in_title = doc.add_paragraph(); run_in_title = p_in_title.add_run("Cash Inflows")
        run_in_title.font.name = 'Helvetica'; run_in_title.bold = True; run_in_title.font.size = Pt(10)
        p_in_title.paragraph_format.space_before = title_para_fmt["space_before"]; p_in_title.paragraph_format.space_after = title_para_fmt["space_after"]
        inflow_items_list = [
            ("Monthly Dues Collected", self.variables['monthly_dues']), ("Certifications Issued", self.variables['certifications']),
            ("Membership Fee", self.variables['membership_fee']), ("Vehicle Stickers", self.variables['vehicle_stickers']),
            ("Rentals", self.variables['rentals']), ("Solicitations/Donations", self.variables['solicitations']),
            ("Interest Income on Bank Deposits", self.variables['interest_income']), 
            ("Livelihood Management Fee", self.variables['livelihood_fee']),
            ("Withdrawal from Bank", self.variables['withdrawal_from_bank']), 
            ("Others (Inflow)", self.variables['inflows_others']), 
            ("Total Cash Receipts", self.variables['total_receipts'])]
        in_table = doc.add_table(rows=len(inflow_items_list), cols=2); in_table.style = 'Table Grid'; in_table.autofit = False
        in_table.columns[0].width = table_col1_width; in_table.columns[1].width = table_col2_width
        for i, (label, var) in enumerate(inflow_items_list):
            is_total_row = (label == "Total Cash Receipts")
            set_cell_style(in_table.cell(i,0), label, bold=is_total_row); set_cell_style(in_table.cell(i,1), format_amount_docx(var), align='right', bold=is_total_row)

        p_out_title = doc.add_paragraph(); run_out_title = p_out_title.add_run("Less: Cash Outflows")
        run_out_title.font.name = 'Helvetica'; run_out_title.bold = True; run_out_title.font.size = Pt(10)
        p_out_title.paragraph_format.space_before = title_para_fmt["space_before"]; p_out_title.paragraph_format.space_after = title_para_fmt["space_after"]
        outflow_items_list = [
            ("Snacks/Meals for Visitors", self.variables['snacks_meals']), ("Transportation Expenses", self.variables['transportation']),
            ("Office Supplies Expense", self.variables['office_supplies']), ("Printing and Photocopy", self.variables['printing']),
            ("Labor", self.variables['labor']), ("Billboard Expense", self.variables['billboard']),
            ("Clearing/Cleaning Charges", self.variables['cleaning']), ("Miscellaneous Expenses", self.variables['misc_expenses']),
            ("Federation Fee", self.variables['federation_fee']), ("HOA-BOD Uniforms", self.variables['uniforms']),
            ("BOD Meeting", self.variables['bod_mtg']), ("General Assembly", self.variables['general_assembly']),
            ("Cash Deposit to Bank", self.variables['cash_deposit']), 
            ("Withholding Tax on Bank Deposit", self.variables['withholding_tax']),
            ("Refund", self.variables['refund']), ("Others (Outflow)", self.variables['outflows_others']),
            ("Total Cash Outflows/Disbursements", self.variables['cash_outflows'])]
        out_table = doc.add_table(rows=len(outflow_items_list), cols=2); out_table.style = 'Table Grid'; out_table.autofit = False
        out_table.columns[0].width = table_col1_width; out_table.columns[1].width = table_col2_width
        for i, (label, var) in enumerate(outflow_items_list):
            is_total_row = (label == "Total Cash Outflows/Disbursements")
            set_cell_style(out_table.cell(i,0), label, bold=is_total_row); set_cell_style(out_table.cell(i,1), format_amount_docx(var), align='right', bold=is_total_row)

        p_end_title = doc.add_paragraph(); run_end_title = p_end_title.add_run("Ending Cash Balance")
        run_end_title.font.name = 'Helvetica'; run_end_title.bold = True; run_end_title.font.size = Pt(10)
        p_end_title.paragraph_format.space_before = title_para_fmt["space_before"]; p_end_title.paragraph_format.space_after = title_para_fmt["space_after"]
        end_table = doc.add_table(rows=1, cols=2); end_table.style = 'Table Grid'; end_table.autofit = False
        end_table.columns[0].width = table_col1_width; end_table.columns[1].width = table_col2_width
        set_cell_style(end_table.cell(0,0), "Ending Cash Balance", bold=True); set_cell_style(end_table.cell(0,1), format_amount_docx(self.variables['ending_cash']), align='right', bold=True)

        p_brk_title = doc.add_paragraph(); run_brk_title = p_brk_title.add_run("Breakdown of Cash")
        run_brk_title.font.name = 'Helvetica'; run_brk_title.bold = True; run_brk_title.font.size = Pt(10)
        p_brk_title.paragraph_format.space_before = title_para_fmt["space_before"]; p_brk_title.paragraph_format.space_after = title_para_fmt["space_after"]
        brk_table = doc.add_table(rows=2, cols=2); brk_table.style = 'Table Grid'; brk_table.autofit = False
        brk_table.columns[0].width = table_col1_width; brk_table.columns[1].width = table_col2_width
        set_cell_style(brk_table.cell(0,0), "Cash in Bank"); set_cell_style(brk_table.cell(0,1), format_amount_docx(self.variables['ending_cash_bank']), align='right')
        set_cell_style(brk_table.cell(1,0), "Cash on Hand"); set_cell_style(brk_table.cell(1,1), format_amount_docx(self.variables['ending_cash_hand']), align='right')
        
        doc.add_paragraph().paragraph_format.space_before = Pt(12) 

        prepared_name = self.prepared_by_var.get() or "_______________________"
        noted_name_1 = self.noted_by_var_1.get() or "_______________________"
        noted_name_2 = self.noted_by_var_2.get() or "_______________________"
        checked_name = self.checked_by_var.get() or "_______________________"
        # Ensure calculated widths are integers (EMUs)
        sign_col_width_val = int((table_content_width / 2) - Inches(0.05))

        prep_check_table = doc.add_table(rows=3, cols=2); prep_check_table.autofit = False
        prep_check_table.columns[0].width = sign_col_width_val; prep_check_table.columns[1].width = sign_col_width_val
        prep_check_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        def set_signatory_name_cell(cell, name_text):
            tc_el = cell._tc
            for p_el in tc_el.xpath('./w:p'): tc_el.remove(p_el)
            para = cell.add_paragraph()
            run = para.add_run("\n" + name_text) 
            run.font.name = 'Helvetica'; run.font.size = Pt(8); run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        set_cell_style(prep_check_table.cell(0,0), "Prepared by:", align='center', size=8)
        set_signatory_name_cell(prep_check_table.cell(1,0), prepared_name)
        set_cell_style(prep_check_table.cell(2,0), "HOA Treasurer", align='center', size=8)
        
        set_cell_style(prep_check_table.cell(0,1), "Checked by:", align='center', size=8)
        set_signatory_name_cell(prep_check_table.cell(1,1), checked_name)
        set_cell_style(prep_check_table.cell(2,1), "HOA Auditor", align='center', size=8)
        
        p_noted_title = doc.add_paragraph(); run_noted_title = p_noted_title.add_run("Noted by:")
        run_noted_title.font.name = 'Helvetica'; run_noted_title.font.size = Pt(8)
        p_noted_title.alignment = WD_ALIGN_PARAGRAPH.CENTER; p_noted_title.paragraph_format.space_before = Pt(10); p_noted_title.paragraph_format.space_after = Pt(4)

        noted_table = doc.add_table(rows=2, cols=2); noted_table.autofit = False
        noted_table.columns[0].width = sign_col_width_val; noted_table.columns[1].width = sign_col_width_val
        noted_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        set_signatory_name_cell(noted_table.cell(0,0), noted_name_1)
        set_cell_style(noted_table.cell(1,0), "HOA President", align='center', size=8)
        
        set_signatory_name_cell(noted_table.cell(0,1), noted_name_2)
        set_cell_style(noted_table.cell(1,1), "CHUDD HCD-CORDS", align='center', size=8)
        
        return doc

    def _create_docx_at_path(self, filename):
        try:
            logging.debug(f"Starting to build DOCX object for {filename}")
            doc = self._build_docx_document()
            logging.debug(f"DOCX object built. Attempting to save to {filename}")
            doc.save(filename)
            logging.info(f"DOCX saved successfully to {filename}")
            return True
        except Exception as e:
            logging.exception(f"Error creating DOCX content for {filename}. THIS IS THE DETAILED TRACEBACK:")
            return False

    def save_to_docx(self):
        try:
            default_filename = f"cash_flow_statement_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word documents", "*.docx")],
                initialfile=default_filename,
                title="Save Word Document As"
            )
            if not filename:
                return {"status": "cancelled", "message": "Save to Word cancelled by user."}

            if self._create_docx_at_path(filename):
                logging.info(f"Word document successfully saved to {filename}")
                return {"status": "success", "message": f"Word document saved to {filename}", "filename": filename}
            else:
                # This message is shown in the GUI. The actual error is in the logs.
                return {"status": "error", "message": f"Failed to create Word document at {filename}.\nCheck logs and ensure python-docx is installed."}
        except Exception as e:
            logging.exception("Error during Word save process") # Outer exception handler
            return {"status": "error", "message": f"An unexpected error occurred during Word save: {str(e)}"}

    def generate_temp_docx(self):
        try:
            fd, temp_filename = tempfile.mkstemp(suffix=".docx", prefix="cash_flow_")
            os.close(fd)
            if self._create_docx_at_path(temp_filename):
                logging.info(f"Temporary DOCX for email generated at {temp_filename}")
                return {"status": "success", "filename": temp_filename}
            else:
                if os.path.exists(temp_filename):
                    try: os.remove(temp_filename)
                    except Exception as e_del: logging.error(f"Could not remove temp DOCX {temp_filename} after generation failure: {e_del}")
                return {"status": "error", "message": "Failed to generate temporary DOCX."}
        except Exception as e:
            logging.exception("Error generating temporary DOCX")
            return {"status": "error", "message": f"An unexpected error occurred while generating temporary DOCX: {str(e)}"}

    def _clear_app_variables(self):
        """Clears data variables in the application, keeping config like logo/address."""
        fields_to_keep = {
            'logo_path_var', 'address_var', 'date_var', 'display_date', 'title_var',
            'recipient_emails_var', 
            'prepared_by_var', 'noted_by_var_1', 'noted_by_var_2', 'checked_by_var',
            'footer_image1_var', 'footer_image2_var', 'calculator' 
        }
        for key, var_obj in self.variables.items():
            if key not in fields_to_keep and hasattr(var_obj, 'set'): 
                var_obj.set("")
        logging.info("Application data variables cleared for loading new file.")


    def load_from_docx(self, filename):
        try:
            logging.info(f"Loading DOCX: {filename}")
            doc = Document(filename)
            self._clear_app_variables()

            logo_extracted = False
            for section in doc.sections:
                header = section.header
                if header and not logo_extracted:
                    for rel_id, rel in header.part.rels.items():
                        if "image" in rel.target_ref: 
                            image_part = rel.target_part; image_data = image_part.blob
                            image_ext = image_part.content_type.split("/")[-1]
                            if image_ext == "octet-stream": image_ext = "png" 
                            with tempfile.NamedTemporaryFile(suffix=f".{image_ext}", prefix="logo_docx_", delete=False) as temp_logo_file:
                                temp_logo_file.write(image_data); image_filename = temp_logo_file.name
                            self.variables['logo_path_var'].set(image_filename); logo_extracted = True; break 
                if logo_extracted: break

            address_found = False; date_found = False
            for section in doc.sections:
                header = section.header
                if header:
                    for table_obj in header.tables: # Renamed to avoid conflict with ReportLab Table
                        if len(table_obj.rows) > 0 and len(table_obj.columns) > 1:
                            text_cell_content = table_obj.cell(0,1).text.strip()
                            lines = [line.strip() for line in text_cell_content.splitlines() if line.strip()]
                            if not address_found and lines and "CASH FLOW STATEMENT" not in lines[0].upper():
                                self.address_var.set(lines[0]); address_found = True
                            if not date_found:
                                date_match = re.search(r"For the Month of\s+(\w+\s+\d{1,2},\s+\d{4})", text_cell_content, re.IGNORECASE)
                                if date_match:
                                    try:
                                        date_obj = datetime.datetime.strptime(date_match.group(1).strip(), "%B %d, %Y")
                                        self.date_var.set(date_obj.strftime("%m/%d/%Y")); date_found = True
                                    except ValueError: logging.warning(f"Could not parse date from DOCX header: {date_match.group(1)}")
                            if address_found and date_found: break
                    if address_found and date_found: break
                if address_found and date_found: break
            if not address_found: self.address_var.set("Default Address - Change Me")

            label_to_var_map = {
                "Cash in Bank-beg": self.variables['cash_bank_beg'], "Cash on Hand-beg": self.variables['cash_hand_beg'],
                "Monthly Dues Collected": self.variables['monthly_dues'], "Certifications Issued": self.variables['certifications'],
                "Membership Fee": self.variables['membership_fee'], "Vehicle Stickers": self.variables['vehicle_stickers'],
                "Rentals": self.variables['rentals'], "Solicitations/Donations": self.variables['solicitations'],
                "Interest Income on Bank Deposits": self.variables['interest_income'], "Livelihood Management Fee": self.variables['livelihood_fee'],
                "Withdrawal from Bank": self.variables['withdrawal_from_bank'], "Others (Inflow)": self.variables['inflows_others'],
                "Snacks/Meals for Visitors": self.variables['snacks_meals'], "Transportation Expenses": self.variables['transportation'],
                "Office Supplies Expense": self.variables['office_supplies'], "Printing and Photocopy": self.variables['printing'],
                "Labor": self.variables['labor'], "Billboard Expense": self.variables['billboard'],
                "Clearing/Cleaning Charges": self.variables['cleaning'], "Miscellaneous Expenses": self.variables['misc_expenses'],
                "Federation Fee": self.variables['federation_fee'], "HOA-BOD Uniforms": self.variables['uniforms'],
                "BOD Meeting": self.variables['bod_mtg'], "General Assembly": self.variables['general_assembly'],
                "Cash Deposit to Bank": self.variables['cash_deposit'], "Withholding Tax on Bank Deposit": self.variables['withholding_tax'],
                "Refund": self.variables['refund'], "Others (Outflow)": self.variables['outflows_others'],
            }

            for table_obj in doc.tables: # docx.Document.tables
                for row_idx, row in enumerate(table_obj.rows):
                    if len(row.cells) >= 2:
                        label = row.cells[0].text.strip()
                        value_text = row.cells[1].text.strip()
                        if label in label_to_var_map:
                            label_to_var_map[label].set(self.parse_amount(value_text))
                        
                        if "Prepared by:" in label and row_idx + 2 < len(table_obj.rows): 
                             name_text = table_obj.cell(row_idx + 1, 0).text.strip()
                             if name_text and "HOA Treasurer" not in table_obj.cell(row_idx+2, 0).text: self.prepared_by_var.set(name_text.splitlines()[0])
                        if "Checked by:" in label and row_idx + 2 < len(table_obj.rows) and len(table_obj.columns) > 1:
                             name_text = table_obj.cell(row_idx + 1, 1).text.strip()
                             if name_text and "HOA Auditor" not in table_obj.cell(row_idx+2,1).text: self.checked_by_var.set(name_text.splitlines()[0])
            
            noted_by_table_found = False
            # Iterate through body child elements (paragraphs and tables)
            for i, body_child_element in enumerate(doc.element.body):
                if body_child_element.tag.endswith('p'): 
                    current_para_text = ""
                    # Find the high-level Paragraph object corresponding to this XML element
                    for p_obj in doc.paragraphs:
                        if p_obj._element is body_child_element: # Check for identity
                            current_para_text = p_obj.text
                            break
                    
                    if "noted by:" in current_para_text.lower():
                        if i + 1 < len(doc.element.body):
                            next_element = doc.element.body[i+1]
                            if next_element.tag.endswith('tbl'): 
                                noted_table_doc = DocxTable(next_element, doc) 
                                if len(noted_table_doc.rows) >= 2 and len(noted_table_doc.columns) == 2: # Check rows >= 2 for safety
                                    # Ensure cell text extraction is safe for potentially empty/merged cells
                                    cell_1_0_text = noted_table_doc.cell(1,0).text.lower() if len(noted_table_doc.rows) > 1 else ""
                                    cell_1_1_text = noted_table_doc.cell(1,1).text.lower() if len(noted_table_doc.rows) > 1 else ""
                                    
                                    if "hoa president" in cell_1_0_text and \
                                       "chudd" in cell_1_1_text:
                                        name1_lines = noted_table_doc.cell(0,0).text.strip().splitlines()
                                        name2_lines = noted_table_doc.cell(0,1).text.strip().splitlines()
                                        if name1_lines: self.noted_by_var_1.set(name1_lines[0])
                                        if name2_lines: self.noted_by_var_2.set(name2_lines[0])
                                        noted_by_table_found = True; break
            if not noted_by_table_found: logging.warning("DOCX Load: 'Noted by' signatory table not found with expected structure.")

            messagebox.showinfo("Success", "DOCX data loaded successfully.")
            return True

        except Exception as e:
            logging.exception(f"Error loading Word document: {filename}")
            messagebox.showerror("Error", f"Error loading Word document:\n{str(e)}")
            return False

    def load_from_pdf(self, filename):
        try:
            logging.info(f"Attempting to load PDF: {filename}")
            self._clear_app_variables()
            address_found = False; date_found = False; logo_extracted = False

            label_to_var_map_pdf = {
                re.sub(r'[^\w\s-]', '', key.lower()): var 
                for key, var in {
                    "Cash in Bank-beg": self.variables['cash_bank_beg'], "Cash on Hand-beg": self.variables['cash_hand_beg'],
                    "Monthly Dues Collected": self.variables['monthly_dues'], "Certifications Issued": self.variables['certifications'],
                    "Membership Fee": self.variables['membership_fee'], "Vehicle Stickers": self.variables['vehicle_stickers'],
                    "Rentals": self.variables['rentals'], "Solicitations/Donations": self.variables['solicitations'],
                    "Interest Income on Bank Deposits": self.variables['interest_income'], "Livelihood Management Fee": self.variables['livelihood_fee'],
                    "Withdrawal from Bank": self.variables['withdrawal_from_bank'], "Others (Inflow)": self.variables['inflows_others'],
                    "Snacks/Meals for Visitors": self.variables['snacks_meals'], "Transportation Expenses": self.variables['transportation'],
                    "Office Supplies Expense": self.variables['office_supplies'], "Printing and Photocopy": self.variables['printing'],
                    "Labor": self.variables['labor'], "Billboard Expense": self.variables['billboard'],
                    "Clearing/Cleaning Charges": self.variables['cleaning'], "Miscellaneous Expenses": self.variables['misc_expenses'],
                    "Federation Fee": self.variables['federation_fee'], "HOA-BOD Uniforms": self.variables['uniforms'], 
                    "BOD Meeting": self.variables['bod_mtg'], "General Assembly": self.variables['general_assembly'],
                    "Cash Deposit to Bank": self.variables['cash_deposit'], "Withholding Tax on Bank Deposit": self.variables['withholding_tax'],
                    "Refund": self.variables['refund'], "Others (Outflow)": self.variables['outflows_others'],
                }.items()
            }
            
            with pdfplumber.open(filename) as pdf:
                first_page = pdf.pages[0]
                page_text_content = first_page.extract_text(x_tolerance=2, y_tolerance=2) 
                if page_text_content:
                    lines = [line.strip() for line in page_text_content.splitlines() if line.strip()]
                    if not address_found:
                        for i, line in enumerate(lines):
                            if "CASH FLOW STATEMENT" in line.upper():
                                if i > 0 and "FOR THE MONTH OF" not in lines[i-1].upper() and not lines[i-1].startswith("[Logo"):
                                    self.address_var.set(lines[i-1]); address_found = True; break
                    if not date_found:
                        date_match_pdf = re.search(r"For the Month of\s+(\w+\s+\d{1,2},\s+\d{4})", page_text_content, re.IGNORECASE)
                        if date_match_pdf:
                            try:
                                date_obj_pdf = datetime.datetime.strptime(date_match_pdf.group(1).strip(), "%B %d, %Y")
                                self.date_var.set(date_obj_pdf.strftime("%m/%d/%Y")); date_found = True
                            except ValueError: logging.warning(f"Could not parse date from PDF text: {date_match_pdf.group(1)}")
                if not address_found: self.address_var.set("Default Address - Change Me")

                for page_num, page in enumerate(pdf.pages):
                    tables_extracted = page.extract_tables(table_settings={"vertical_strategy": "lines", "horizontal_strategy": "text", "snap_tolerance": 5, "intersection_x_tolerance": 3, "text_x_tolerance":3}) 
                    if tables_extracted: # Ensure tables_extracted is not None
                        for table_idx, table_data in enumerate(tables_extracted):
                            if not table_data: continue
                            for row_idx, row in enumerate(table_data):
                                if not row or len(row) < 2 or not row[0]: continue
                                label_raw = str(row[0]).replace('\n', ' ').strip()
                                label_normalized = re.sub(r'[^\w\s-]', '', label_raw.lower())
                                value_raw = str(row[1]).replace('\n', ' ').strip() if len(row) > 1 and row[1] else ""
                                if label_normalized in label_to_var_map_pdf:
                                    label_to_var_map_pdf[label_normalized].set(self.parse_amount(value_raw))
                
                full_text_for_names = "".join([p.extract_text(x_tolerance=1, y_tolerance=1, layout=True) or "" for p in pdf.pages]) 
                
                if not self.prepared_by_var.get():
                    prepared_match = re.search(r"Prepared by:\s*([\s\S]*?)\s*HOA Treasurer", full_text_for_names, re.IGNORECASE)
                    if prepared_match: 
                        name_lines = prepared_match.group(1).strip().splitlines()
                        if name_lines: self.prepared_by_var.set(name_lines[-1].strip()) 
                if not self.checked_by_var.get():
                    checked_match = re.search(r"Checked by:\s*([\s\S]*?)\s*HOA Auditor", full_text_for_names, re.IGNORECASE)
                    if checked_match: 
                        name_lines = checked_match.group(1).strip().splitlines()
                        if name_lines: self.checked_by_var.set(name_lines[-1].strip())
                
                if not self.noted_by_var_1.get() or not self.noted_by_var_2.get():
                    noted_by_block_match = re.search(r"Noted by:\s*([\s\S]*?)(?=Prepared by:|Checked by:|$)", full_text_for_names, re.IGNORECASE) 
                    if noted_by_block_match:
                        noted_block_text = noted_by_block_match.group(1)
                        if not self.noted_by_var_1.get():
                            noted1_m = re.search(r"([\w\s.]+?)\s*HOA President", noted_block_text, re.IGNORECASE | re.DOTALL)
                            if noted1_m: 
                                name_lines = noted1_m.group(1).strip().splitlines()
                                if name_lines: self.noted_by_var_1.set(name_lines[-1].strip())
                        if not self.noted_by_var_2.get():
                            noted2_m = re.search(r"([\w\s.]+?)\s*CHUDD HCD-CORDS", noted_block_text, re.IGNORECASE | re.DOTALL)
                            if noted2_m: 
                                name_lines = noted2_m.group(1).strip().splitlines()
                                if name_lines: self.noted_by_var_2.set(name_lines[-1].strip())
            
            if not logo_extracted:
                doc_fitz = fitz.open(filename)
                for page_num in range(min(1, len(doc_fitz))): 
                    page_fitz = doc_fitz[page_num]; images = page_fitz.get_images(full=True)
                    if images: 
                        page_height = page_fitz.rect.height
                        for img_info in images:
                            xref = img_info[0]; bbox = page_fitz.get_image_bbox(img_info) 
                            if bbox.y1 < page_height * 0.30 and bbox.x0 < page_fitz.rect.width * 0.30: 
                                base_image = doc_fitz.extract_image(xref)
                                image_bytes = base_image["image"]; image_ext = base_image["ext"]
                                if image_ext == "jpx": image_ext = "jpeg" 
                                elif image_ext == "jpeg2000": image_ext = "jp2"
                                with tempfile.NamedTemporaryFile(suffix=f".{image_ext}", prefix="logo_pdf_", delete=False) as temp_logo_file:
                                    temp_logo_file.write(image_bytes); image_filename = temp_logo_file.name
                                self.variables['logo_path_var'].set(image_filename); logo_extracted = True; break 
                    if logo_extracted: break
                doc_fitz.close()
            
            messagebox.showinfo("Success", "PDF data loaded successfully!")
            return True
        except Exception as e:
            logging.exception(f"Error loading PDF: {filename}")
            messagebox.showerror("Error", f"Error loading PDF:\n{str(e)}")
            return False
        
    def load_from_documentpdf(self):
        try:
            filename = filedialog.askopenfilename(
                filetypes=[("Documents", "*.docx *.pdf"), ("Word Documents", "*.docx"), ("PDF Files", "*.pdf"), ("All Files", "*.*")],
                title="Select a Document (DOCX or PDF)"
            )
            if not filename: return

            if filename.lower().endswith('.pdf'): self.load_from_pdf(filename)
            elif filename.lower().endswith('.docx'): self.load_from_docx(filename)
            else: messagebox.showerror("Error", "Unsupported file format. Please select a PDF or DOCX file.")
            
            if self.calculator and hasattr(self.calculator, 'calculate_totals'):
                self.calculator.calculate_totals()
            else:
                logging.warning("Calculator object not found or 'calculate_totals' method missing after load.")

        except Exception as e:
             logging.exception("Error loading document")
             messagebox.showerror("Error", f"Error loading document: {str(e)}")