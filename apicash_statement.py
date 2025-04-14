import os
import datetime
import smtplib
import sys
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import csv
from decimal import Decimal
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.shared import Pt, Inches
import customtkinter as ctk
from tkcalendar import Calendar

# Add HoverCalendar class from the first code
class HoverCalendar(Calendar):
    """Custom Calendar class to enable hovering over month and year for navigation."""
    def __init__(self, master=None, **kw):
        super().__init__(master, **kw)
        self._setup_hover_navigation()

    def _setup_hover_navigation(self):
        """Add hover bindings to month and year labels."""
        self._header_month_label = self._calendar[0][2]  # Month label widget
        self._header_month_label.bind("<Enter>", self._on_month_hover)
        self._header_month_label.bind("<Leave>", self._on_month_leave)
        self._header_year_label = self._calendar[0][4]  # Year label widget
        self._header_year_label.bind("<Enter>", self._on_year_hover)
        self._header_year_label.bind("<Leave>", self._on_year_leave)

    def _on_month_hover(self, event):
        self._header_month_label.configure(fg="blue")
        self._calendar[0][3].event_generate("<Button-1>")

    def _on_month_leave(self, event):
        self._header_month_label.configure(fg=self["foreground"])

    def _on_year_hover(self, event):
        self._header_year_label.configure(fg="blue")
        self._date = self._date.replace(year=self._date.year + 1)
        self._setup_calendar()

    def _on_year_leave(self, event):
        self._header_year_label.configure(fg=self["foreground"])

class IntegratedCashFlowApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Cash Flow Statement Generator with Email")
        self.root.geometry("800x700")  # Set initial size to 1200x1000
        
        # Set CustomTkinter appearance (from first code)
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")

        # Color scheme (from first code)
        self.primary_color = "#1C2526"  # Dark background
        self.secondary_color = "#2A3F4D"  # Section backgrounds
        self.accent_color = "#00A7E1"  # Buttons
        self.text_color = "#E0E0E0"  # Light grey text
        self.success_color = "#4CAF50"  # Success indicators
        
        # Hardcoded email credentials
        self.SENDER_EMAIL = "chuddcdo@gmail.com"
        self.SENDER_PASSWORD = "jfyb eoog ukxr hhiq"
        
        # Recipient email variable
        self.recipient_emails_var = ctk.StringVar()
        
        # Cash flow variables
        self.title_var = ctk.StringVar(value="Statement Of Cash Flows")
        self.date_var = ctk.StringVar(value=datetime.datetime.now().strftime("%m/%d/%Y"))
        self.display_date = ctk.StringVar(value=datetime.datetime.now().strftime("%b %d, %Y"))
        
        self.cash_bank_beg = ctk.StringVar()
        self.cash_hand_beg = ctk.StringVar()
        self.monthly_dues = ctk.StringVar()
        self.certifications = ctk.StringVar()
        self.membership_fee = ctk.StringVar()
        self.vehicle_stickers = ctk.StringVar()
        self.rentals = ctk.StringVar()
        self.solicitations = ctk.StringVar()
        self.interest_income = ctk.StringVar()
        self.livelihood_fee = ctk.StringVar()
        self.inflows_others = ctk.StringVar()
        self.total_receipts = ctk.StringVar()
        self.cash_outflows = ctk.StringVar()
        self.snacks_meals = ctk.StringVar()
        self.transportation = ctk.StringVar()
        self.office_supplies = ctk.StringVar()
        self.printing = ctk.StringVar()
        self.labor = ctk.StringVar()
        self.billboard = ctk.StringVar()
        self.cleaning = ctk.StringVar()
        self.misc_expenses = ctk.StringVar()
        self.federation_fee = ctk.StringVar()
        self.uniforms = ctk.StringVar()
        self.bod_mtg = ctk.StringVar()
        self.general_assembly = ctk.StringVar()
        self.cash_deposit = ctk.StringVar()
        self.withholding_tax = ctk.StringVar()
        self.refund_sericulture = ctk.StringVar()
        self.outflows_others = ctk.StringVar()
        self.ending_cash = ctk.StringVar()
        self.ending_cash_bank = ctk.StringVar()
        self.ending_cash_hand = ctk.StringVar()
        
        # Split ratios for ending cash balances
        self.bank_split_ratio = Decimal('0.8')  # 80% to bank
        self.hand_split_ratio = Decimal('0.2')  # 20% to hand
        
        self.input_vars = [
            self.cash_bank_beg, self.cash_hand_beg, self.monthly_dues, self.certifications,
            self.membership_fee, self.vehicle_stickers, self.rentals, self.solicitations,
            self.interest_income, self.livelihood_fee, self.inflows_others, self.snacks_meals,
            self.transportation, self.office_supplies, self.printing, self.labor, self.billboard,
            self.cleaning, self.misc_expenses, self.federation_fee, self.uniforms, self.bod_mtg,
            self.general_assembly, self.cash_deposit, self.withholding_tax, self.refund_sericulture,
            self.outflows_others
        ]
        
        for var in self.input_vars:
            var.trace_add('write', lambda *args: self.calculate_totals())
        
        # Sync display_date with date_var
        self.date_var.trace('w', self.update_display_date)
        
        self.create_widgets()
        self.setup_keyboard_shortcuts()

    def update_display_date(self, *args):
        """Convert mm/dd/yyyy from date_var to MMM dd, yyyy for display_date."""
        raw_date = self.date_var.get()
        try:
            date_obj = datetime.datetime.strptime(raw_date, "%m/%d/%Y")
            self.display_date.set(date_obj.strftime("%b %d, %Y"))
        except ValueError:
            pass

    def setup_keyboard_shortcuts(self):
        self.root.bind('<Control-s>', lambda e: self.save_to_csv())
        self.root.bind('<Control-e>', lambda e: self.export_to_pdf())
        self.root.bind('<Control-l>', lambda e: self.load_from_csv())
        self.root.bind('<Control-g>', lambda e: self.send_email())
        self.root.bind('<Control-w>', lambda e: self.save_to_docx())

    def format_entry(self, var, entry_widget):
        def on_change(*args):
            value = var.get()
            if value:
                try:
                    formatted = f"{Decimal(value.replace(',', '')):,.2f}"
                    if formatted != value:
                        var.set(formatted)
                except:
                    pass
        var.trace_add('write', on_change)
        entry_widget.configure(justify="right")

    def validate_date(self, date_str):
        if not date_str:
            return True
        try:
            datetime.datetime.strptime(date_str, "%m/%d/%Y")
            return True
        except ValueError:
            messagebox.showwarning("Invalid Date", "Please enter date in mm/dd/yyyy format.")
            return False

    def create_tooltip(self, widget, text):
        tooltip = tk.Toplevel(widget)
        tooltip.wm_overrideredirect(True)
        tooltip.wm_geometry("+1000+1000")
        label = tk.Label(tooltip, text=text, background="lightyellow", relief="solid", borderwidth=1, fg="black")
        label.pack()
        
        def show(event):
            x = widget.winfo_rootx() + 20
            y = widget.winfo_rooty() + 20
            tooltip.wm_geometry(f"+{x}+{y}")
            tooltip.deiconify()
        
        def hide(event):
            tooltip.withdraw()
        
        widget.bind("<Enter>", show)
        widget.bind("<Leave>", hide)
        tooltip.withdraw()

    def show_calendar(self):
        """Show a standalone calendar in a popup window, appearing directly in the center."""
        popup = tk.Toplevel(self.root)
        popup.title("Select Date")
        popup.geometry("300x300")
        popup.transient(self.root)
        popup.grab_set()
        popup.withdraw()

        popup.update_idletasks()
        popup_width = 300
        popup_height = 300
        main_width = self.root.winfo_width()
        main_height = self.root.winfo_height()
        main_x = self.root.winfo_x()
        main_y = self.root.winfo_y()

        x = main_x + (main_width - popup_width) // 2
        y = main_y + (main_height - popup_height) // 2
        popup.geometry(f"{popup_width}x{popup_height}+{x}+{y}")

        cal = HoverCalendar(
            popup,
            font=("Arial", 12),
            background="#2A3F4D",
            foreground="#E0E0E0",
            selectbackground="#2A3F4D",
            selectforeground="#E0E0E0",
            normalbackground="#FFFFFF",
            normalforeground="#000000",
            weekendbackground="#FFFFFF",
            weekendforeground="#000000",
            headersbackground="#2A3F4D",
            headersforeground="#E0E0E0",
            showothermonthdays=False,
            showweeknumbers=False
        )
        cal.pack(padx=10, pady=10)

        try:
            current_date = datetime.datetime.strptime(self.date_var.get(), "%m/%d/%Y")
            cal.selection_set(current_date)
        except ValueError:
            pass

        def on_date_select():
            selected_date = cal.selection_get()
            if selected_date:
                self.date_var.set(selected_date.strftime("%m/%d/%Y"))
            popup.destroy()

        confirm_button = ctk.CTkButton(
            popup,
            text="Select",
            command=on_date_select,
            fg_color=self.accent_color,
            text_color=self.text_color,
            hover_color="#008CC1"
        )
        confirm_button.pack(pady=10)

        popup.deiconify()
        popup.protocol("WM_DELETE_WINDOW", popup.destroy)

    def create_widgets(self):
        self.main_frame = ctk.CTkFrame(self.root, fg_color=self.primary_color)
        self.main_frame.pack(fill="both", expand=True, padx=10, pady=10)  # Added margins for breathing space
        
        # Removed canvas and scrollbar setup since scrolling is not needed
        self.scrollable_frame = ctk.CTkFrame(self.main_frame, fg_color=self.primary_color)
        self.scrollable_frame.pack(fill="both", expand=True, padx=10, pady=5)  # Added internal margins
        
        # Title and Date Frame
        header_frame = ctk.CTkFrame(self.scrollable_frame, fg_color=self.secondary_color, corner_radius=10)
        header_frame.grid(row=0, column=0, sticky="ew", padx=10, pady=5)  # Added margins
        
        ctk.CTkLabel(header_frame, text="Title:", font=("Roboto", 12), text_color=self.text_color).pack(side="left", padx=5)
        title_entry = ctk.CTkEntry(header_frame, textvariable=self.title_var, width=200, font=("Roboto", 12), fg_color="#3A4F5D", text_color=self.text_color)
        title_entry.pack(side="left", padx=5)
        
        date_frame = ctk.CTkFrame(header_frame, fg_color="transparent")
        date_frame.pack(side="left", padx=5)
        ctk.CTkLabel(date_frame, text="Date:", font=("Roboto", 12), text_color=self.text_color).pack(side="left")
        date_button = ctk.CTkButton(
            date_frame,
            textvariable=self.display_date,
            font=("Arial", 12),
            fg_color="#3A4F5D",
            text_color=self.text_color,
            command=self.show_calendar,
            width=100,
            height=28,
            corner_radius=5,
            hover_color="#4A5F6D"
        )
        date_button.pack(side="left")
        self.create_tooltip(date_button, "Click to select a date from the calendar")
        
        # Email Configuration Frame
        email_frame = ctk.CTkFrame(self.scrollable_frame, fg_color=self.secondary_color, corner_radius=10)
        email_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=5)  # Added margins
        
        ctk.CTkLabel(email_frame, text="Recipients (comma-separated):", font=("Roboto", 12), text_color=self.text_color).pack(side="left", padx=5)
        email_entry = ctk.CTkEntry(email_frame, textvariable=self.recipient_emails_var, width=300, font=("Roboto", 12), fg_color="#3A4F5D", text_color=self.text_color)
        email_entry.pack(side="left", padx=5)
        
        # Buttons Frame
        button_frame = ctk.CTkFrame(self.scrollable_frame, fg_color=self.primary_color)
        button_frame.grid(row=2, column=0, sticky="ew", padx=10, pady=5)  # Added margins
        
        buttons = [
            ("Save to CSV (Ctrl+S)", self.save_to_csv),
            ("Load from CSV (Ctrl+L)", self.load_from_csv),
            ("Clear All Fields", self.clear_fields),
            ("Export to PDF (Ctrl+E)", self.export_to_pdf),
            ("Save to Word (Ctrl+W)", self.save_to_docx),
            ("Send via Email (Ctrl+G)", self.send_email),
        ]
        for text, command in buttons:
            ctk.CTkButton(
                button_frame,
                text=text,
                command=command,
                font=("Roboto", 12),
                fg_color=self.accent_color,
                hover_color="#008CC1",
                text_color=self.text_color,
                width=150,
                height=35,
                corner_radius=8
            ).pack(side="left", padx=5)
        
        # Columns Frame for responsive layout
        self.columns_frame = ctk.CTkFrame(self.scrollable_frame, fg_color=self.primary_color)
        self.columns_frame.grid(row=3, column=0, sticky="nsew", padx=10, pady=5)  # Added margins
        
        # Define column frames
        self.beg_frame = ctk.CTkFrame(self.columns_frame, fg_color=self.secondary_color, corner_radius=10)
        ctk.CTkLabel(self.beg_frame, text="Beginning Cash Balances", font=("Roboto", 14, "bold"), text_color=self.text_color).pack(anchor="w", padx=5, pady=5)

        self.inflow_frame = ctk.CTkFrame(self.columns_frame, fg_color=self.secondary_color, corner_radius=10)
        ctk.CTkLabel(self.inflow_frame, text="Cash Inflows", font=("Roboto", 14, "bold"), text_color=self.text_color).pack(anchor="w", padx=5, pady=5)

        self.outflow_frame = ctk.CTkFrame(self.columns_frame, fg_color=self.secondary_color, corner_radius=10)
        ctk.CTkLabel(self.outflow_frame, text="Cash Outflows", font=("Roboto", 14, "bold"), text_color=self.text_color).pack(anchor="w", padx=5, pady=5)

        self.end_frame = ctk.CTkFrame(self.columns_frame, fg_color=self.secondary_color, corner_radius=10)
        ctk.CTkLabel(self.end_frame, text="Ending Cash Balances", font=("Roboto", 14, "bold"), text_color=self.text_color).pack(anchor="w", padx=5, pady=5)

        self.totals_frame = ctk.CTkFrame(self.columns_frame, fg_color=self.secondary_color, corner_radius=10)
        ctk.CTkLabel(self.totals_frame, text="Totals", font=("Roboto", 14, "bold"), text_color=self.text_color).pack(anchor="w", padx=5, pady=5)
        
        # Store frames for layout management
        self.column_frames = [
            self.beg_frame,
            self.inflow_frame,
            self.outflow_frame,
            self.end_frame,
            self.totals_frame
        ]
        
        # Populate column frames
        self.populate_columns()
        
        # Bind resize event
        self.root.bind("<Configure>", self.update_layout)

    def populate_columns(self):
        # Beginning Balances
        beg_inner = ctk.CTkFrame(self.beg_frame, fg_color="transparent")
        beg_inner.pack(fill="x", padx=5, pady=5)
        beg_items = [
            ("Cash in Bank (beginning):", self.cash_bank_beg),
            ("Cash on Hand (beginning):", self.cash_hand_beg)
        ]
        for i, (label, var) in enumerate(beg_items):
            ctk.CTkLabel(beg_inner, text=label, font=("Roboto", 11), text_color=self.text_color, anchor="w").grid(row=i, column=0, sticky="w", padx=5, pady=2)
            entry = ctk.CTkEntry(beg_inner, textvariable=var, width=120, font=("Roboto", 11), fg_color="#3A4F5D", text_color=self.text_color)
            entry.grid(row=i, column=1, sticky="e", padx=5, pady=2)
            self.format_entry(var, entry)
        
        # Cash Inflows
        inflow_inner = ctk.CTkFrame(self.inflow_frame, fg_color="transparent")
        inflow_inner.pack(fill="x", padx=5, pady=5)
        inflow_items = [
            ("Monthly dues collected:", self.monthly_dues),
            ("Certifications issued:", self.certifications),
            ("Membership fee:", self.membership_fee),
            ("Vehicle stickers:", self.vehicle_stickers),
            ("Rentals (covered courts):", self.rentals),
            ("Solicitations/Donations:", self.solicitations),
            ("Interest Income:", self.interest_income),
            ("Livelihood Management Fee:", self.livelihood_fee),
            ("Others:", self.inflows_others)
        ]
        for i, (label, var) in enumerate(inflow_items):
            ctk.CTkLabel(inflow_inner, text=label, font=("Roboto", 11), text_color=self.text_color, anchor="w").grid(row=i, column=0, sticky="w", padx=5, pady=2)
            entry = ctk.CTkEntry(inflow_inner, textvariable=var, width=120, font=("Roboto", 11), fg_color="#3A4F5D", text_color=self.text_color)
            entry.grid(row=i, column=1, sticky="e", padx=5, pady=2)
            self.format_entry(var, entry)
        
        # Cash Outflows
        outflow_inner = ctk.CTkFrame(self.outflow_frame, fg_color="transparent")
        outflow_inner.pack(fill="x", padx=5, pady=5)
        outflow_items = [
            ("Snacks/Meals for visitors:", self.snacks_meals),
            ("Transportation expenses:", self.transportation),
            ("Office supplies expense:", self.office_supplies),
            ("Printing and photocopy:", self.printing),
            ("Labor:", self.labor),
            ("Billboard expense:", self.billboard),
            ("Clearing/cleaning charges:", self.cleaning),
            ("Miscellaneous expenses:", self.misc_expenses),
            ("Federation fee:", self.federation_fee),
            ("HOA-BOD Uniforms:", self.uniforms),
            ("BOD Mtg:", self.bod_mtg),
            ("General Assembly:", self.general_assembly),
            ("Cash Deposit to bank:", self.cash_deposit),
            ("Withholding tax:", self.withholding_tax),
            ("Refund for seri-culture:", self.refund_sericulture),
            ("Others:", self.outflows_others)
        ]
        for i, (label, var) in enumerate(outflow_items):
            ctk.CTkLabel(outflow_inner, text=label, font=("Roboto", 11), text_color=self.text_color, anchor="w").grid(row=i, column=0, sticky="w", padx=5, pady=2)
            entry = ctk.CTkEntry(outflow_inner, textvariable=var, width=120, font=("Roboto", 11), fg_color="#3A4F5D", text_color=self.text_color)
            entry.grid(row=i, column=1, sticky="e", padx=5, pady=2)
            self.format_entry(var, entry)
        
        # Ending Balances
        end_inner = ctk.CTkFrame(self.end_frame, fg_color="transparent")
        end_inner.pack(fill="x", padx=5, pady=5)
        end_items = [
            ("Cash in Bank:", self.ending_cash_bank),
            ("Cash on Hand:", self.ending_cash_hand)
        ]
        for i, (label, var) in enumerate(end_items):
            ctk.CTkLabel(end_inner, text=label, font=("Roboto", 11), text_color=self.text_color, anchor="w").grid(row=i, column=0, sticky="w", padx=5, pady=2)
            entry = ctk.CTkEntry(end_inner, textvariable=var, width=120, font=("Roboto", 11), fg_color="#3A4F5D", text_color=self.text_color, state="disabled")
            entry.grid(row=i, column=1, sticky="e", padx=5, pady=2)
        
        # Totals
        total_inner = ctk.CTkFrame(self.totals_frame, fg_color="transparent")
        total_inner.pack(fill="x", padx=5, pady=5)
        total_items = [
            ("Total Cash Receipts:", self.total_receipts),
            ("Cash Outflows:", self.cash_outflows),
            ("Ending Cash Balance:", self.ending_cash)
        ]
        for i, (label, var) in enumerate(total_items):
            ctk.CTkLabel(total_inner, text=label, font=("Roboto", 11), text_color=self.text_color, anchor="w").grid(row=i, column=0, sticky="w", padx=5, pady=2)
            entry = ctk.CTkEntry(total_inner, textvariable=var, width=120, font=("Roboto", 11), fg_color="#3A4F5D", text_color=self.text_color, state="disabled")
            entry.grid(row=i, column=1, sticky="e", padx=5, pady=2)

    def update_layout(self, event=None):
        window_width = self.main_frame.winfo_width()
        min_column_width = 250  # Increased to ensure content fits without truncation
        num_columns = max(1, window_width // min_column_width)
        num_columns = min(num_columns, len(self.column_frames))
        
        for frame in self.column_frames:
            frame.grid_forget()
        
        for i, frame in enumerate(self.column_frames):
            row = i // num_columns
            col = i % num_columns
            frame.grid(row=row, column=col, sticky="nsew", padx=5, pady=5)  # Added margins between columns
        
        # Ensure the columns_frame expands to fill available space
        self.columns_frame.grid_columnconfigure(tuple(range(num_columns)), weight=1)
        self.columns_frame.grid_rowconfigure(tuple(range((len(self.column_frames) + num_columns - 1) // num_columns)), weight=1)

    def safe_decimal(self, var):
        val = var.get().strip()
        if not val:
            return Decimal("0")
        try:
            val = val.replace(",", "")
            return Decimal(val)
        except:
            return Decimal("0")

    def calculate_totals(self):
        try:
            inflow_total = sum([
                self.safe_decimal(self.monthly_dues),
                self.safe_decimal(self.certifications),
                self.safe_decimal(self.membership_fee),
                self.safe_decimal(self.vehicle_stickers),
                self.safe_decimal(self.rentals),
                self.safe_decimal(self.solicitations),
                self.safe_decimal(self.interest_income),
                self.safe_decimal(self.livelihood_fee),
                self.safe_decimal(self.inflows_others)
            ])
            
            outflow_total = sum([
                self.safe_decimal(self.snacks_meals),
                self.safe_decimal(self.transportation),
                self.safe_decimal(self.office_supplies),
                self.safe_decimal(self.printing),
                self.safe_decimal(self.labor),
                self.safe_decimal(self.billboard),
                self.safe_decimal(self.cleaning),
                self.safe_decimal(self.misc_expenses),
                self.safe_decimal(self.federation_fee),
                self.safe_decimal(self.uniforms),
                self.safe_decimal(self.bod_mtg),
                self.safe_decimal(self.general_assembly),
                self.safe_decimal(self.cash_deposit),
                self.safe_decimal(self.withholding_tax),
                self.safe_decimal(self.refund_sericulture),
                self.safe_decimal(self.outflows_others)
            ])
            
            beginning_total = self.safe_decimal(self.cash_bank_beg) + self.safe_decimal(self.cash_hand_beg)
            ending_balance = beginning_total + inflow_total - outflow_total
            
            self.total_receipts.set(f"{inflow_total:,.2f}")
            self.cash_outflows.set(f"{outflow_total:,.2f}")
            self.ending_cash.set(f"{ending_balance:,.2f}")
            
            ending_cash_bank = ending_balance * self.bank_split_ratio
            ending_cash_hand = ending_balance * self.hand_split_ratio
            self.ending_cash_bank.set(f"{ending_cash_bank:,.2f}")
            self.ending_cash_hand.set(f"{ending_cash_hand:,.2f}")
            
        except Exception:
            self.total_receipts.set("")
            self.cash_outflows.set("")
            self.ending_cash.set("")
            self.ending_cash_bank.set("")
            self.ending_cash_hand.set("")

    def format_date_for_display(self, date_str):
        try:
            date_obj = datetime.datetime.strptime(date_str, "%m/%d/%Y")
            return date_obj.strftime("%B %d, %Y")
        except ValueError:
            return date_str

    def format_date_for_entry(self, date_str):
        try:
            date_obj = datetime.datetime.strptime(date_str, "%B %d, %Y")
            return date_obj.strftime("%m/%d/%Y")
        except ValueError:
            return date_str

    def save_to_csv(self):
        try:
            self.calculate_totals()
            filename = f"cash_flow_statement_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
            with open(filename, 'w', newline='') as csvfile:
                writer = csv.writer(csvfile)
                writer.writerow([self.title_var.get()])
                writer.writerow([f"For the year month {self.format_date_for_display(self.date_var.get())}"])
                writer.writerow([])
                writer.writerow(["Cash in Bank-beg", self.cash_bank_beg.get()])
                writer.writerow(["Cash on Hand-beg", self.cash_hand_beg.get()])
                writer.writerow([])
                writer.writerow(["Cash inflows:"])
                writer.writerow(["Monthly dues collected", self.monthly_dues.get()])
                writer.writerow(["Certifications issued", self.certifications.get()])
                writer.writerow(["Membership fee", self.membership_fee.get()])
                writer.writerow(["Vehicle stickers", self.vehicle_stickers.get()])
                writer.writerow(["Rentals (covered courts)", self.rentals.get()])
                writer.writerow(["Solicitations/Donations", self.solicitations.get()])
                writer.writerow(["Interest Income on bank deposits", self.interest_income.get()])
                writer.writerow(["Livelihood Management Fee", self.livelihood_fee.get()])
                writer.writerow(["Others", self.inflows_others.get()])
                writer.writerow(["Total Cash receipt", self.total_receipts.get()])
                writer.writerow([])
                writer.writerow(["Less:"])
                writer.writerow(["Cash Out Flows/Disbursements", self.cash_outflows.get()])
                writer.writerow(["Snacks/Meals for visitors", self.snacks_meals.get()])
                writer.writerow(["Transportation expenses", self.transportation.get()])
                writer.writerow(["Office supplies expense", self.office_supplies.get()])
                writer.writerow(["Printing and photocopy", self.printing.get()])
                writer.writerow(["Labor", self.labor.get()])
                writer.writerow(["Billboard expense", self.billboard.get()])
                writer.writerow(["Clearing/cleaning charges", self.cleaning.get()])
                writer.writerow(["Miscellaneous expenses", self.misc_expenses.get()])
                writer.writerow(["Federation fee", self.federation_fee.get()])
                writer.writerow(["HOA-BOD Uniforms", self.uniforms.get()])
                writer.writerow(["BOD Mtg", self.bod_mtg.get()])
                writer.writerow(["General Assembly", self.general_assembly.get()])
                writer.writerow(["Cash Deposit to bank", self.cash_deposit.get()])
                writer.writerow(["Withholding tax on bank deposit", self.withholding_tax.get()])
                writer.writerow(["Refund for seri-culture", self.refund_sericulture.get()])
                writer.writerow(["Others", self.outflows_others.get()])
                writer.writerow([])
                writer.writerow(["Ending cash balance", self.ending_cash.get()])
                writer.writerow([])
                writer.writerow(["Breakdown of cash:"])
                writer.writerow(["Cash in Bank", self.ending_cash_bank.get()])
                writer.writerow(["Cash on Hand", self.ending_cash_hand.get()])
            
            messagebox.showinfo("Success", f"Cash flow statement saved to {filename}")
            return filename
            
        except Exception as e:
            messagebox.showerror("Error", f"Error saving to CSV: {str(e)}")
            return None

    def load_from_csv(self):
        try:
            filename = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv")])
            if not filename:
                return
                
            with open(filename, 'r') as csvfile:
                reader = csv.reader(csvfile)
                data = list(reader)
                self.title_var.set(data[0][0])
                date_str = data[1][0].replace("For the year month ", "")
                self.date_var.set(self.format_date_for_entry(date_str))
                self.cash_bank_beg.set(data[3][1])
                self.cash_hand_beg.set(data[4][1])
                self.monthly_dues.set(data[7][1])
                self.certifications.set(data[8][1])
                self.membership_fee.set(data[9][1])
                self.vehicle_stickers.set(data[10][1])
                self.rentals.set(data[11][1])
                self.solicitations.set(data[12][1])
                self.interest_income.set(data[13][1])
                self.livelihood_fee.set(data[14][1])
                self.inflows_others.set(data[15][1])
                self.total_receipts.set(data[16][1])
                self.cash_outflows.set(data[19][1])
                self.snacks_meals.set(data[20][1])
                self.transportation.set(data[21][1])
                self.office_supplies.set(data[22][1])
                self.printing.set(data[23][1])
                self.labor.set(data[24][1])
                self.billboard.set(data[25][1])
                self.cleaning.set(data[26][1])
                self.misc_expenses.set(data[27][1])
                self.federation_fee.set(data[28][1])
                self.uniforms.set(data[29][1])
                self.bod_mtg.set(data[30][1])
                self.general_assembly.set(data[31][1])
                self.cash_deposit.set(data[32][1])
                self.withholding_tax.set(data[33][1])
                self.refund_sericulture.set(data[34][1])
                self.outflows_others.set(data[35][1])
                self.ending_cash.set(data[37][1])
                self.ending_cash_bank.set(data[40][1])
                self.ending_cash_hand.set(data[41][1])
            
            messagebox.showinfo("Success", f"Loaded data from {filename}")
            self.calculate_totals()
            
        except Exception as e:
            messagebox.showerror("Error", f"Error loading CSV: {str(e)}")

    def clear_fields(self):
        for var in self.input_vars + [self.ending_cash_bank, self.ending_cash_hand]:
            var.set("")
        
        self.total_receipts.set("")
        self.cash_outflows.set("")
        self.ending_cash.set("")
        
        messagebox.showinfo("Success", "All fields have been cleared")

    def export_to_pdf(self):
        try:
            self.calculate_totals()

            def format_amount(value):
                if value:
                    try:
                        amount = Decimal(value.replace(',', ''))
                        return f"{amount:,.2f}"
                    except:
                        return value
                return ""

            default_filename = f"cash_flow_statement_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            filename = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF files", "*.pdf")],
                initialfile=default_filename,
                title="Save PDF As"
            )
            if not filename:
                return None

            doc = SimpleDocTemplate(filename, pagesize=letter)
            styles = getSampleStyleSheet()
            elements = []

            if getattr(sys, 'frozen', False):
                base_path = sys._MEIPASS
            else:
                base_path = os.path.dirname(__file__)

            logo_path = os.path.join(base_path, "logo.png")

            if os.path.exists(logo_path):
                elements.append(Image(logo_path, width=100, height=100))
                elements.append(Spacer(1, 12))

            elements.append(Paragraph(self.title_var.get(), styles['Title']))
            elements.append(Paragraph(f"For the year month {self.format_date_for_display(self.date_var.get())}", styles['Normal']))
            elements.append(Spacer(1, 12))
            
            beg_data = [
                ["Cash in Bank-beg", format_amount(self.cash_bank_beg.get())],
                ["Cash on Hand-beg", format_amount(self.cash_hand_beg.get())]
            ]
            beg_table = Table(beg_data, colWidths=[300, 150])
            beg_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ]))
            elements.append(beg_table)
            elements.append(Spacer(1, 12))
            
            elements.append(Paragraph("<b>Cash inflows:</b>", styles['Normal']))
            elements.append(Spacer(1, 6))
            inflows_data = [
                ["Monthly dues collected", format_amount(self.monthly_dues.get())],
                ["Certifications issued", format_amount(self.certifications.get())],
                ["Membership fee", format_amount(self.membership_fee.get())],
                ["Vehicle stickers", format_amount(self.vehicle_stickers.get())],
                ["Rentals (covered courts)", format_amount(self.rentals.get())],
                ["Solicitations/Donations", format_amount(self.solicitations.get())],
                ["Interest Income on bank deposits", format_amount(self.interest_income.get())],
                ["Livelihood Management Fee", format_amount(self.livelihood_fee.get())],
                ["Others", format_amount(self.inflows_others.get())],
                ["Total Cash receipt", format_amount(self.total_receipts.get())]
            ]
            inflows_table = Table(inflows_data, colWidths=[300, 150])
            inflows_table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('BACKGROUND', (0, -1), (-1, -1), colors.lightgrey),
                ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
                ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ]))
            elements.append(inflows_table)
            elements.append(Spacer(1, 12))
            
            elements.append(Paragraph("<b>Less:</b>", styles['Normal']))
            elements.append(Spacer(1, 6))
            outflows_data = [
                ["Cash Out Flows/Disbursements", format_amount(self.cash_outflows.get())],
                ["Snacks/Meals for visitors", format_amount(self.snacks_meals.get())],
                ["Transportation expenses", format_amount(self.transportation.get())],
                ["Office supplies expense", format_amount(self.office_supplies.get())],
                ["Printing and photocopy", format_amount(self.printing.get())],
                ["Labor", format_amount(self.labor.get())],
                ["Billboard expense", format_amount(self.billboard.get())],
                ["Clearing/cleaning charges", format_amount(self.cleaning.get())],
                ["Miscellaneous expenses", format_amount(self.misc_expenses.get())],
                ["Federation fee", format_amount(self.federation_fee.get())],
                ["HOA-BOD Uniforms", format_amount(self.uniforms.get())],
                ["BOD Mtg", format_amount(self.bod_mtg.get())],
                ["General Assembly", format_amount(self.general_assembly.get())],
                ["Cash Deposit to bank", format_amount(self.cash_deposit.get())],
                ["Withholding tax on bank deposit", format_amount(self.withholding_tax.get())],
                ["Refund for seri-culture", format_amount(self.refund_sericulture.get())],
                ["Others", format_amount(self.outflows_others.get())]
            ]
            outflows_table = Table(outflows_data, colWidths=[300, 150])
            outflows_table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('BACKGROUND', (0, 0), (0, 0), colors.lightgrey),
                ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
                ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ]))
            elements.append(outflows_table)
            elements.append(Spacer(1, 12))
            
            ending_data = [
                ["Ending cash balance", format_amount(self.ending_cash.get())]
            ]
            ending_table = Table(ending_data, colWidths=[300, 150])
            ending_table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ]))
            elements.append(ending_table)
            elements.append(Spacer(1, 12))
            
            elements.append(Paragraph("<b>Breakdown of cash:</b>", styles['Normal']))
            elements.append(Spacer(1, 6))
            breakdown_data = [
                ["Cash in Bank", format_amount(self.ending_cash_bank.get())],
                ["Cash on Hand", format_amount(self.ending_cash_hand.get())]
            ]
            breakdown_table = Table(breakdown_data, colWidths=[300, 150])
            breakdown_table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ]))
            elements.append(breakdown_table)
            
            def add_page_numbers(canvas, doc):
                page_num = canvas.getPageNumber()
                text = f"Page {page_num}"
                canvas.drawRightString(200, 20, text)
            
            doc.build(elements, onFirstPage=add_page_numbers, onLaterPages=add_page_numbers)
            messagebox.showinfo("Success", f"PDF successfully exported to {filename}")
            return filename
            
        except Exception as e:
            messagebox.showerror("Error", f"Error exporting to PDF: {str(e)}\n\nMake sure you have ReportLab installed by running:\npip install reportlab")
            return None

    def save_to_docx(self):
        try:
            self.calculate_totals()

            def format_amount(value):
                if value:
                    try:
                        amount = Decimal(value.replace(',', ''))
                        return f"{amount:,.2f}"
                    except:
                        return value
                return ""

            default_filename = f"cash_flow_statement_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word documents", "*.docx")],
                initialfile=default_filename,
                title="Save Word Document As"
            )
            if not filename:
                return None

            doc = Document()

            if getattr(sys, 'frozen', False):
                base_path = sys._MEIPASS
            else:
                base_path = os.path.dirname(__file__)

            logo_path = os.path.join(base_path, "logo.png")

            if os.path.exists(logo_path):
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(logo_path, width=Inches(1.5))
                paragraph.alignment = 1

            doc.add_heading(self.title_var.get(), level=1)
            doc.add_paragraph(f"For the year month {self.format_date_for_display(self.date_var.get())}")

            doc.add_heading("Beginning Cash Balances", level=2)
            table = doc.add_table(rows=2, cols=2)
            table.style = 'Table Grid'
            table.cell(0, 0).text = "Cash in Bank-beg"
            table.cell(0, 1).text = format_amount(self.cash_bank_beg.get())
            table.cell(1, 0).text = "Cash on Hand-beg"
            table.cell(1, 1).text = format_amount(self.cash_hand_beg.get())
            for row in table.rows:
                row.cells[1].paragraphs[0].alignment = 2

            doc.add_heading("Cash Inflows", level=2)
            table = doc.add_table(rows=10, cols=2)
            table.style = 'Table Grid'
            inflow_items = [
                ("Monthly dues collected", self.monthly_dues),
                ("Certifications issued", self.certifications),
                ("Membership fee", self.membership_fee),
                ("Vehicle stickers", self.vehicle_stickers),
                ("Rentals (covered courts)", self.rentals),
                ("Solicitations/Donations", self.solicitations),
                ("Interest Income on bank deposits", self.interest_income),
                ("Livelihood Management Fee", self.livelihood_fee),
                ("Others", self.inflows_others),
                ("Total Cash receipt", self.total_receipts)
            ]
            for i, (label, var) in enumerate(inflow_items):
                table.cell(i, 0).text = label
                table.cell(i, 1).text = format_amount(var.get())
                table.cell(i, 1).paragraphs[0].alignment = 2
            table.cell(9, 0).paragraphs[0].runs[0].bold = True

            doc.add_heading("Less: Cash Outflows", level=2)
            table = doc.add_table(rows=17, cols=2)
            table.style = 'Table Grid'
            outflow_items = [
                ("Cash Out Flows/Disbursements", self.cash_outflows),
                ("Snacks/Meals for visitors", self.snacks_meals),
                ("Transportation expenses", self.transportation),
                ("Office supplies expense", self.office_supplies),
                ("Printing and photocopy", self.printing),
                ("Labor", self.labor),
                ("Billboard expense", self.billboard),
                ("Clearing/cleaning charges", self.cleaning),
                ("Miscellaneous expenses", self.misc_expenses),
                ("Federation fee", self.federation_fee),
                ("HOA-BOD Uniforms", self.uniforms),
                ("BOD Mtg", self.bod_mtg),
                ("General Assembly", self.general_assembly),
                ("Cash Deposit to bank", self.cash_deposit),
                ("Withholding tax on bank deposit", self.withholding_tax),
                ("Refund for seri-culture", self.refund_sericulture),
                ("Others", self.outflows_others)
            ]
            for i, (label, var) in enumerate(outflow_items):
                table.cell(i, 0).text = label
                table.cell(i, 1).text = format_amount(var.get())
                table.cell(i, 1).paragraphs[0].alignment = 2
            table.cell(0, 0).paragraphs[0].runs[0].bold = True

            doc.add_heading("Ending Cash Balance", level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.cell(0, 0).text = "Ending cash balance"
            table.cell(0, 1).text = format_amount(self.ending_cash.get())
            table.cell(0, 1).paragraphs[0].alignment = 2
            table.cell(0, 0).paragraphs[0].runs[0].bold = True

            doc.add_heading("Breakdown of Cash", level=2)
            table = doc.add_table(rows=2, cols=2)
            table.style = 'Table Grid'
            table.cell(0, 0).text = "Cash in Bank"
            table.cell(0, 1).text = format_amount(self.ending_cash_bank.get())
            table.cell(1, 0).text = "Cash on Hand"
            table.cell(1, 1).text = format_amount(self.ending_cash_hand.get())
            for row in table.rows:
                row.cells[1].paragraphs[0].alignment = 2

            doc.save(filename)
            messagebox.showinfo("Success", f"Word document saved to {filename}")
            return filename
            
        except Exception as e:
            messagebox.showerror("Error", f"Error saving to Word: {str(e)}\n\nMake sure you have python-docx installed by running:\npip install python-docx")
            return None

    def send_email(self):
        try:
            sender_email = self.SENDER_EMAIL
            sender_password = self.SENDER_PASSWORD
            recipient_emails = [email.strip() for email in self.recipient_emails_var.get().split(',') if email.strip()]
            
            if not recipient_emails:
                messagebox.showerror("Error", "Please fill in the recipient email field.")
                return

            pdf_filename = self.export_to_pdf()
            if not pdf_filename:
                return
            docx_filename = self.save_to_docx()
            if not docx_filename:
                os.remove(pdf_filename)
                return

            msg = MIMEMultipart()
            msg['From'] = sender_email
            msg['To'] = ", ".join(recipient_emails)
            msg['Subject'] = f"Cash Flow Statement - {self.format_date_for_display(self.date_var.get())}"

            body = f"Attached is the cash flow statement for {self.format_date_for_display(self.date_var.get())} in both PDF and Word formats.\n\nRegards,\nYour's truly"
            msg.attach(MIMEText(body, 'plain'))

            with open(pdf_filename, 'rb') as f:
                part = MIMEApplication(f.read(), Name=os.path.basename(pdf_filename))
                part['Content-Disposition'] = f'attachment; filename="{os.path.basename(pdf_filename)}"'
                msg.attach(part)

            with open(docx_filename, 'rb') as f:
                part = MIMEApplication(f.read(), Name=os.path.basename(docx_filename))
                part['Content-Disposition'] = f'attachment; filename="{os.path.basename(docx_filename)}"'
                msg.attach(part)

            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login(sender_email, sender_password)
            
            server.send_message(msg)
            server.quit()
            
            messagebox.showinfo("Success", f"Email with PDF and Word files sent to {', '.join(recipient_emails)}!")
            os.remove(pdf_filename)
            os.remove(docx_filename)
            
        except smtplib.SMTPAuthenticationError as e:
            messagebox.showerror("Error", f"Authentication failed: {str(e)}\nCheck your hardcoded email and app password.")
            if 'pdf_filename' in locals():
                os.remove(pdf_filename)
            if 'docx_filename' in locals():
                os.remove(docx_filename)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to send email: {str(e)}\nEnsure your email credentials are correct and you have an internet connection.")
            if 'pdf_filename' in locals():
                os.remove(pdf_filename)
            if 'docx_filename' in locals():
                os.remove(docx_filename)

if __name__ == "__main__":
    root = ctk.CTk()  # Use ctk.CTk instead of tk.Tk
    app = IntegratedCashFlowApp(root)
    root.mainloop()
