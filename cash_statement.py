import os
import datetime
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import tkinter as tk
from tkinter import ttk, messagebox
import csv
from decimal import Decimal
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

# Integrated Cash Flow Statement App
class IntegratedCashFlowApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Cash Flow Statement Generator with Email")
        self.root.geometry("800x800")
        
        # Email variables (to be filled by user in GUI)
        self.sender_email_var = tk.StringVar()
        self.sender_password_var = tk.StringVar()
        self.recipient_emails_var = tk.StringVar()
        
        # Cash flow variables
        self.title_var = tk.StringVar(value="Statement Of Cash Flows")
        self.today_date = datetime.datetime.now().strftime("%B %d, %Y")
        
        self.cash_bank_beg = tk.StringVar()
        self.cash_hand_beg = tk.StringVar()
        
        self.monthly_dues = tk.StringVar()
        self.certifications = tk.StringVar()
        self.membership_fee = tk.StringVar()
        self.vehicle_stickers = tk.StringVar()
        self.rentals = tk.StringVar()
        self.solicitations = tk.StringVar()
        self.interest_income = tk.StringVar()
        self.livelihood_fee = tk.StringVar()
        self.inflows_others = tk.StringVar()
        self.total_receipts = tk.StringVar()
        
        self.cash_outflows = tk.StringVar()
        self.snacks_meals = tk.StringVar()
        self.transportation = tk.StringVar()
        self.office_supplies = tk.StringVar()
        self.printing = tk.StringVar()
        self.labor = tk.StringVar()
        self.billboard = tk.StringVar()
        self.cleaning = tk.StringVar()
        self.misc_expenses = tk.StringVar()
        self.federation_fee = tk.StringVar()
        self.uniforms = tk.StringVar()
        self.bod_mtg = tk.StringVar()
        self.general_assembly = tk.StringVar()
        self.cash_deposit = tk.StringVar()
        self.withholding_tax = tk.StringVar()
        self.refund_sericulture = tk.StringVar()
        self.outflows_others = tk.StringVar()
        self.outflows_others_2 = tk.StringVar()
        
        self.ending_cash = tk.StringVar()
        self.ending_cash_bank = tk.StringVar()
        self.ending_cash_hand = tk.StringVar()
        
        self.create_widgets()
        self.setup_keyboard_shortcuts()

    def setup_keyboard_shortcuts(self):
        self.root.bind('<Control-s>', lambda e: self.save_to_csv())
        self.root.bind('<Control-e>', lambda e: self.export_to_pdf())
        self.root.bind('<Control-l>', lambda e: self.load_from_csv())
        self.root.bind('<Control-g>', lambda e: self.send_email())
        self.root.bind('<Control-w>', lambda e: self.save_to_docx())

    def format_entry(self, var, entry_widget):
        def on_change(*args):
            value = var.get()
            if value:
                try:
                    formatted = f"{Decimal(value.replace(',', '')):,.2f}"
                    var.set(formatted)
                except:
                    pass
        var.trace('w', on_change)
        entry_widget.config(justify='right')

    def create_widgets(self):
        main_frame = ttk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        canvas = tk.Canvas(main_frame)
        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        # Title, Date, and Email Fields
        title_frame = ttk.Frame(scrollable_frame)
        title_frame.pack(fill="x", pady=5)
        
        ttk.Label(title_frame, text="Title:").pack(side="left", padx=5)
        ttk.Entry(title_frame, textvariable=self.title_var, width=40).pack(side="left", padx=5)
        ttk.Label(title_frame, text=f"Date: {self.today_date}").pack(side="left", padx=20)
        
        # Email Configuration Frame
        email_frame = ttk.LabelFrame(scrollable_frame, text="Email Configuration")
        email_frame.pack(fill="x", pady=5)
        
        ttk.Label(email_frame, text="Sender Email:").grid(row=0, column=0, sticky="w", padx=5, pady=2)
        ttk.Entry(email_frame, textvariable=self.sender_email_var, width=30).grid(row=0, column=1, padx=5, pady=2)
        
        ttk.Label(email_frame, text="App Password:").grid(row=1, column=0, sticky="w", padx=5, pady=2)
        ttk.Entry(email_frame, textvariable=self.sender_password_var, width=30, show="*").grid(row=1, column=1, padx=5, pady=2)
        
        ttk.Label(email_frame, text="Recipients (comma-separated):").grid(row=2, column=0, sticky="w", padx=5, pady=2)
        ttk.Entry(email_frame, textvariable=self.recipient_emails_var, width=30).grid(row=2, column=1, padx=5, pady=2)
        
        # Beginning Cash Balances
        beg_frame = ttk.LabelFrame(scrollable_frame, text="Beginning Cash Balances")
        beg_frame.pack(fill="x", pady=5)
        
        ttk.Label(beg_frame, text="Cash in Bank (beginning):").grid(row=0, column=0, sticky="w", padx=5, pady=2)
        bank_beg_entry = ttk.Entry(beg_frame, textvariable=self.cash_bank_beg, width=15)
        bank_beg_entry.grid(row=0, column=1, padx=5, pady=2)
        self.format_entry(self.cash_bank_beg, bank_beg_entry)
        
        ttk.Label(beg_frame, text="Cash on Hand (beginning):").grid(row=1, column=0, sticky="w", padx=5, pady=2)
        hand_beg_entry = ttk.Entry(beg_frame, textvariable=self.cash_hand_beg, width=15)
        hand_beg_entry.grid(row=1, column=1, padx=5, pady=2)
        self.format_entry(self.cash_hand_beg, hand_beg_entry)
        
        # Cash Inflows
        inflow_frame = ttk.LabelFrame(scrollable_frame, text="Cash Inflows")
        inflow_frame.pack(fill="x", pady=5)
        
        inflow_items = [
            ("Monthly dues collected:", self.monthly_dues),
            ("Certifications issued:", self.certifications),
            ("Membership fee:", self.membership_fee),
            ("Vehicle stickers:", self.vehicle_stickers),
            ("Rentals (covered courts):", self.rentals),
            ("Solicitations/Donations:", self.solicitations),
            ("Interest Income on bank deposits:", self.interest_income),
            ("Livelihood Management Fee:", self.livelihood_fee),
            ("Others:", self.inflows_others),
        ]
        
        for i, (label, var) in enumerate(inflow_items):
            ttk.Label(inflow_frame, text=label).grid(row=i, column=0, sticky="w", padx=5, pady=2)
            entry = ttk.Entry(inflow_frame, textvariable=var, width=15)
            entry.grid(row=i, column=1, padx=5, pady=2)
            self.format_entry(var, entry)
        
        ttk.Label(inflow_frame, text="Total Cash receipt:").grid(row=len(inflow_items), column=0, sticky="w", padx=5, pady=2)
        ttk.Entry(inflow_frame, textvariable=self.total_receipts, width=15, state="readonly").grid(row=len(inflow_items), column=1, padx=5, pady=2)
        
        # Cash Outflows
        outflow_frame = ttk.LabelFrame(scrollable_frame, text="Cash Outflows")
        outflow_frame.pack(fill="x", pady=5)
        
        ttk.Label(outflow_frame, text="Cash Out Flows/Disbursements:").grid(row=0, column=0, sticky="w", padx=5, pady=2)
        ttk.Entry(outflow_frame, textvariable=self.cash_outflows, width=15, state="readonly").grid(row=0, column=1, padx=5, pady=2)
        
        outflow_items = [
            ("Snacks/Meals for visitors:", self.snacks_meals),
            ("Transportation expenses:", self.transportation),
            ("Office supplies expense:", self.office_supplies),
            ("Printing and photocopy:", self.printing),
            ("Labor:", self.labor),
            ("Billboard expense:", self.billboard),
            ("Clearing/cleaning charges:", self.cleaning),
            ("Miscellaneous expenses:", self.misc_expenses),
            ("Federation fee:", self.federation_fee),
            ("HOA-BOD Uniforms:", self.uniforms),
            ("BOD Mtg:", self.bod_mtg),
            ("General Assembly:", self.general_assembly),
            ("Cash Deposit to bank:", self.cash_deposit),
            ("Withholding tax on bank deposit:", self.withholding_tax),
            ("Refund for seri-culture:", self.refund_sericulture),
            ("Others:", self.outflows_others),
            ("", self.outflows_others_2)
        ]
        
        for i, (label, var) in enumerate(outflow_items):
            ttk.Label(outflow_frame, text=label).grid(row=i+1, column=0, sticky="w", padx=5, pady=2)
            entry = ttk.Entry(outflow_frame, textvariable=var, width=15)
            entry.grid(row=i+1, column=1, padx=5, pady=2)
            self.format_entry(var, entry)
        
        # Ending Cash Balances
        end_frame = ttk.LabelFrame(scrollable_frame, text="Ending Cash Balances")
        end_frame.pack(fill="x", pady=5)
        
        ttk.Label(end_frame, text="Ending cash balance:").grid(row=0, column=0, sticky="w", padx=5, pady=2)
        ttk.Entry(end_frame, textvariable=self.ending_cash, width=15, state="readonly").grid(row=0, column=1, padx=5, pady=2)
        
        ttk.Label(end_frame, text="Cash in Bank:").grid(row=1, column=0, sticky="w", padx=5, pady=2)
        bank_end_entry = ttk.Entry(end_frame, textvariable=self.ending_cash_bank, width=15)
        bank_end_entry.grid(row=1, column=1, padx=5, pady=2)
        self.format_entry(self.ending_cash_bank, bank_end_entry)
        
        ttk.Label(end_frame, text="Cash on Hand:").grid(row=2, column=0, sticky="w", padx=5, pady=2)
        hand_end_entry = ttk.Entry(end_frame, textvariable=self.ending_cash_hand, width=15)
        hand_end_entry.grid(row=2, column=1, padx=5, pady=2)
        self.format_entry(self.ending_cash_hand, hand_end_entry)
        
        # Buttons Frame
        button_frame = ttk.Frame(scrollable_frame)
        button_frame.pack(fill="x", pady=10)
        
        ttk.Button(button_frame, text="Save to CSV (Ctrl+S)", command=self.save_to_csv).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Load from CSV (Ctrl+L)", command=self.load_from_csv).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Clear All Fields", command=self.clear_fields).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Export to PDF (Ctrl+E)", command=self.export_to_pdf).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Save to Word (Ctrl+W)", command=self.save_to_docx).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Send via Email (Ctrl+G)", command=self.send_email).pack(side="left", padx=5)

    def safe_decimal(self, var):
        val = var.get().strip()
        if not val:
            return Decimal("0")
        try:
            val = val.replace(",", "")
            return Decimal(val)
        except:
            raise ValueError(f"Invalid number: {val}")

    def calculate_totals(self):
        try:
            inflow_total = sum([
                self.safe_decimal(self.monthly_dues),
                self.safe_decimal(self.certifications),
                self.safe_decimal(self.membership_fee),
                self.safe_decimal(self.vehicle_stickers),
                self.safe_decimal(self.rentals),
                self.safe_decimal(self.solicitations),
                self.safe_decimal(self.interest_income),
                self.safe_decimal(self.livelihood_fee),
                self.safe_decimal(self.inflows_others)
            ])
            
            outflow_total = sum([
                self.safe_decimal(self.snacks_meals),
                self.safe_decimal(self.transportation),
                self.safe_decimal(self.office_supplies),
                self.safe_decimal(self.printing),
                self.safe_decimal(self.labor),
                self.safe_decimal(self.billboard),
                self.safe_decimal(self.cleaning),
                self.safe_decimal(self.misc_expenses),
                self.safe_decimal(self.federation_fee),
                self.safe_decimal(self.uniforms),
                self.safe_decimal(self.bod_mtg),
                self.safe_decimal(self.general_assembly),
                self.safe_decimal(self.cash_deposit),
                self.safe_decimal(self.withholding_tax),
                self.safe_decimal(self.refund_sericulture),
                self.safe_decimal(self.outflows_others),
                self.safe_decimal(self.outflows_others_2)
            ])
            
            beginning_total = self.safe_decimal(self.cash_bank_beg) + self.safe_decimal(self.cash_hand_beg)
            ending_balance = beginning_total + inflow_total - outflow_total
            
            self.total_receipts.set(f"{inflow_total:,.2f}")
            self.cash_outflows.set(f"{outflow_total:,.2f}")
            self.ending_cash.set(f"{ending_balance:,.2f}")
            
            if not self.ending_cash_bank.get() and not self.ending_cash_hand.get():
                self.ending_cash_bank.set(f"{ending_balance * Decimal('0.8'):,.2f}")
                self.ending_cash_hand.set(f"{ending_balance * Decimal('0.2'):,.2f}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Calculation error: {str(e)}\nPlease ensure all values are valid numbers.")

    def save_to_csv(self):
        try:
            filename = f"cash_flow_statement_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
            with open(filename, 'w', newline='') as csvfile:
                writer = csv.writer(csvfile)
                writer.writerow([self.title_var.get()])
                writer.writerow([f"For the year month {self.today_date}"])
                writer.writerow([])
                writer.writerow(["Cash in Bank-beg", self.cash_bank_beg.get()])
                writer.writerow(["Cash on Hand-beg", self.cash_hand_beg.get()])
                writer.writerow([])
                writer.writerow(["Cash inflows:"])
                writer.writerow(["Monthly dues collected", self.monthly_dues.get()])
                writer.writerow(["Certifications issued", self.certifications.get()])
                writer.writerow(["Membership fee", self.membership_fee.get()])
                writer.writerow(["Vehicle stickers", self.vehicle_stickers.get()])
                writer.writerow(["Rentals (covered courts)", self.rentals.get()])
                writer.writerow(["Solicitations/Donations", self.solicitations.get()])
                writer.writerow(["Interest Income on bank deposits", self.interest_income.get()])
                writer.writerow(["Livelihood Management Fee", self.livelihood_fee.get()])
                writer.writerow(["Others", self.inflows_others.get()])
                writer.writerow(["Total Cash receipt", self.total_receipts.get()])
                writer.writerow([])
                writer.writerow(["Less:"])
                writer.writerow(["Cash Out Flows/Disbursements", self.cash_outflows.get()])
                writer.writerow(["Snacks/Meals for visitors", self.snacks_meals.get()])
                writer.writerow(["Transportation expenses", self.transportation.get()])
                writer.writerow(["Office supplies expense", self.office_supplies.get()])
                writer.writerow(["Printing and photocopy", self.printing.get()])
                writer.writerow(["Labor", self.labor.get()])
                writer.writerow(["Billboard expense", self.billboard.get()])
                writer.writerow(["Clearing/cleaning charges", self.cleaning.get()])
                writer.writerow(["Miscellaneous expenses", self.misc_expenses.get()])
                writer.writerow(["Federation fee", self.federation_fee.get()])
                writer.writerow(["HOA-BOD Uniforms", self.uniforms.get()])
                writer.writerow(["BOD Mtg", self.bod_mtg.get()])
                writer.writerow(["General Assembly", self.general_assembly.get()])
                writer.writerow(["Cash Deposit to bank", self.cash_deposit.get()])
                writer.writerow(["Withholding tax on bank deposit", self.withholding_tax.get()])
                writer.writerow(["Refund for seri-culture", self.refund_sericulture.get()])
                writer.writerow(["Others", self.outflows_others.get()])
                writer.writerow(["", self.outflows_others_2.get()])
                writer.writerow([])
                writer.writerow(["Ending cash balance", self.ending_cash.get()])
                writer.writerow([])
                writer.writerow(["Breakdown of cash:"])
                writer.writerow(["Cash in Bank", self.ending_cash_bank.get()])
                writer.writerow(["Cash on Hand", self.ending_cash_hand.get()])
            
            messagebox.showinfo("Success", f"Cash flow statement saved to {filename}")
            return filename
            
        except Exception as e:
            messagebox.showerror("Error", f"Error saving to CSV: {str(e)}")
            return None

    def load_from_csv(self):
        try:
            from tkinter import filedialog
            filename = filedialog.askopenfilename(filetypes=[("CSV files", "*.csv")])
            if not filename:
                return
                
            with open(filename, 'r') as csvfile:
                reader = csv.reader(csvfile)
                data = list(reader)
                self.title_var.set(data[0][0])
                self.cash_bank_beg.set(data[3][1])
                self.cash_hand_beg.set(data[4][1])
                self.monthly_dues.set(data[7][1])
                self.certifications.set(data[8][1])
                self.membership_fee.set(data[9][1])
                self.vehicle_stickers.set(data[10][1])
                self.rentals.set(data[11][1])
                self.solicitations.set(data[12][1])
                self.interest_income.set(data[13][1])
                self.livelihood_fee.set(data[14][1])
                self.inflows_others.set(data[15][1])
                self.total_receipts.set(data[16][1])
                self.cash_outflows.set(data[19][1])
                self.snacks_meals.set(data[20][1])
                self.transportation.set(data[21][1])
                self.office_supplies.set(data[22][1])
                self.printing.set(data[23][1])
                self.labor.set(data[24][1])
                self.billboard.set(data[25][1])
                self.cleaning.set(data[26][1])
                self.misc_expenses.set(data[27][1])
                self.federation_fee.set(data[28][1])
                self.uniforms.set(data[29][1])
                self.bod_mtg.set(data[30][1])
                self.general_assembly.set(data[31][1])
                self.cash_deposit.set(data[32][1])
                self.withholding_tax.set(data[33][1])
                self.refund_sericulture.set(data[34][1])
                self.outflows_others.set(data[35][1])
                self.outflows_others_2.set(data[36][1])
                self.ending_cash.set(data[38][1])
                self.ending_cash_bank.set(data[41][1])
                self.ending_cash_hand.set(data[42][1])
            
            messagebox.showinfo("Success", f"Loaded data from {filename}")
            
        except Exception as e:
            messagebox.showerror("Error", f"Error loading CSV: {str(e)}")

    def clear_fields(self):
        for var in [
            self.cash_bank_beg, self.cash_hand_beg,
            self.monthly_dues, self.certifications, self.membership_fee,
            self.vehicle_stickers, self.rentals, self.solicitations,
            self.interest_income, self.livelihood_fee, self.inflows_others,
            self.snacks_meals, self.transportation, self.office_supplies,
            self.printing, self.labor, self.billboard, self.cleaning,
            self.misc_expenses, self.federation_fee, self.uniforms,
            self.bod_mtg, self.general_assembly, self.cash_deposit,
            self.withholding_tax, self.refund_sericulture, self.outflows_others,
            self.outflows_others_2, self.ending_cash_bank, self.ending_cash_hand
        ]:
            var.set("")
        
        self.total_receipts.set("")
        self.cash_outflows.set("")
        self.ending_cash.set("")
        
        messagebox.showinfo("Success", "All fields have been cleared")

    def export_to_pdf(self):
        try:
            # Automatically calculate totals before exporting
            self.calculate_totals()

            def format_amount(value):
                if value:
                    try:
                        amount = Decimal(value.replace(',', ''))
                        return f"{amount:,.2f}"
                    except:
                        return value
                return ""

            filename = f"cash_flow_statement_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            doc = SimpleDocTemplate(filename, pagesize=letter)
            styles = getSampleStyleSheet()
            elements = []
            
            if os.path.exists("logo.png"):
                elements.append(Image("logo.png", width=100, height=100))
                elements.append(Spacer(1, 12))
            
            elements.append(Paragraph(self.title_var.get(), styles['Title']))
            elements.append(Paragraph(f"For the year month {self.today_date}", styles['Normal']))
            elements.append(Spacer(1, 12))
            
            beg_data = [
                ["Cash in Bank-beg", format_amount(self.cash_bank_beg.get())],
                ["Cash on Hand-beg", format_amount(self.cash_hand_beg.get())]
            ]
            beg_table = Table(beg_data, colWidths=[300, 150])
            beg_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ]))
            elements.append(beg_table)
            elements.append(Spacer(1, 12))
            
            elements.append(Paragraph("<b>Cash inflows:</b>", styles['Normal']))
            elements.append(Spacer(1, 6))
            inflows_data = [
                ["Monthly dues collected", format_amount(self.monthly_dues.get())],
                ["Certifications issued", format_amount(self.certifications.get())],
                ["Membership fee", format_amount(self.membership_fee.get())],
                ["Vehicle stickers", format_amount(self.vehicle_stickers.get())],
                ["Rentals (covered courts)", format_amount(self.rentals.get())],
                ["Solicitations/Donations", format_amount(self.solicitations.get())],
                ["Interest Income on bank deposits", format_amount(self.interest_income.get())],
                ["Livelihood Management Fee", format_amount(self.livelihood_fee.get())],
                ["Others", format_amount(self.inflows_others.get())],
                ["Total Cash receipt", format_amount(self.total_receipts.get())]
            ]
            inflows_table = Table(inflows_data, colWidths=[300, 150])
            inflows_table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('BACKGROUND', (0, -1), (-1, -1), colors.lightgrey),
                ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
                ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ]))
            elements.append(inflows_table)
            elements.append(Spacer(1, 12))
            
            elements.append(Paragraph("<b>Less:</b>", styles['Normal']))
            elements.append(Spacer(1, 6))
            outflows_data = [
                ["Cash Out Flows/Disbursements", format_amount(self.cash_outflows.get())],
                ["Snacks/Meals for visitors", format_amount(self.snacks_meals.get())],
                ["Transportation expenses", format_amount(self.transportation.get())],
                ["Office supplies expense", format_amount(self.office_supplies.get())],
                ["Printing and photocopy", format_amount(self.printing.get())],
                ["Labor", format_amount(self.labor.get())],
                ["Billboard expense", format_amount(self.billboard.get())],
                ["Clearing/cleaning charges", format_amount(self.cleaning.get())],
                ["Miscellaneous expenses", format_amount(self.misc_expenses.get())],
                ["Federation fee", format_amount(self.federation_fee.get())],
                ["HOA-BOD Uniforms", format_amount(self.uniforms.get())],
                ["BOD Mtg", format_amount(self.bod_mtg.get())],
                ["General Assembly", format_amount(self.general_assembly.get())],
                ["Cash Deposit to bank", format_amount(self.cash_deposit.get())],
                ["Withholding tax on bank deposit", format_amount(self.withholding_tax.get())],
                ["Refund for seri-culture", format_amount(self.refund_sericulture.get())],
                ["Others", format_amount(self.outflows_others.get())],
                ["", format_amount(self.outflows_others_2.get())]
            ]
            outflows_table = Table(outflows_data, colWidths=[300, 150])
            outflows_table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('BACKGROUND', (0, 0), (0, 0), colors.lightgrey),
                ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
                ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ]))
            elements.append(outflows_table)
            elements.append(Spacer(1, 12))
            
            ending_data = [
                ["Ending cash balance", format_amount(self.ending_cash.get())]
            ]
            ending_table = Table(ending_data, colWidths=[300, 150])
            ending_table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ]))
            elements.append(ending_table)
            elements.append(Spacer(1, 12))
            
            elements.append(Paragraph("<b>Breakdown of cash:</b>", styles['Normal']))
            elements.append(Spacer(1, 6))
            breakdown_data = [
                ["Cash in Bank", format_amount(self.ending_cash_bank.get())],
                ["Cash on Hand", format_amount(self.ending_cash_hand.get())]
            ]
            breakdown_table = Table(breakdown_data, colWidths=[300, 150])
            breakdown_table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
            ]))
            elements.append(breakdown_table)
            
            def add_page_numbers(canvas, doc):
                page_num = canvas.getPageNumber()
                text = f"Page {page_num}"
                canvas.drawRightString(200, 20, text)
            
            doc.build(elements, onFirstPage=add_page_numbers, onLaterPages=add_page_numbers)
            messagebox.showinfo("Success", f"PDF successfully exported to {filename}")
            return filename
            
        except Exception as e:
            messagebox.showerror("Error", f"Error exporting to PDF: {str(e)}\n\nMake sure you have ReportLab installed by running:\npip install reportlab")
            return None

    def save_to_docx(self):
        try:
            # Automatically calculate totals before saving
            self.calculate_totals()

            def format_amount(value):
                if value:
                    try:
                        amount = Decimal(value.replace(',', ''))
                        return f"{amount:,.2f}"
                    except:
                        return value
                return ""

            filename = f"cash_flow_statement_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc = Document()

            # Title and Date
            doc.add_heading(self.title_var.get(), level=1)
            doc.add_paragraph(f"For the year month {self.today_date}")

            # Beginning Cash Balances
            doc.add_heading("Beginning Cash Balances", level=2)
            table = doc.add_table(rows=2, cols=2)
            table.style = 'Table Grid'
            table.cell(0, 0).text = "Cash in Bank-beg"
            table.cell(0, 1).text = format_amount(self.cash_bank_beg.get())
            table.cell(1, 0).text = "Cash on Hand-beg"
            table.cell(1, 1).text = format_amount(self.cash_hand_beg.get())
            for row in table.rows:
                row.cells[1].paragraphs[0].alignment = 2  # Right align amounts

            # Cash Inflows
            doc.add_heading("Cash Inflows", level=2)
            table = doc.add_table(rows=10, cols=2)
            table.style = 'Table Grid'
            inflow_items = [
                ("Monthly dues collected", self.monthly_dues),
                ("Certifications issued", self.certifications),
                ("Membership fee", self.membership_fee),
                ("Vehicle stickers", self.vehicle_stickers),
                ("Rentals (covered courts)", self.rentals),
                ("Solicitations/Donations", self.solicitations),
                ("Interest Income on bank deposits", self.interest_income),
                ("Livelihood Management Fee", self.livelihood_fee),
                ("Others", self.inflows_others),
                ("Total Cash receipt", self.total_receipts)
            ]
            for i, (label, var) in enumerate(inflow_items):
                table.cell(i, 0).text = label
                table.cell(i, 1).text = format_amount(var.get())
                table.cell(i, 1).paragraphs[0].alignment = 2  # Right align amounts
            table.cell(9, 0).paragraphs[0].runs[0].bold = True  # Bold total

            # Cash Outflows
            doc.add_heading("Less: Cash Outflows", level=2)
            table = doc.add_table(rows=18, cols=2)
            table.style = 'Table Grid'
            outflow_items = [
                ("Cash Out Flows/Disbursements", self.cash_outflows),
                ("Snacks/Meals for visitors", self.snacks_meals),
                ("Transportation expenses", self.transportation),
                ("Office supplies expense", self.office_supplies),
                ("Printing and photocopy", self.printing),
                ("Labor", self.labor),
                ("Billboard expense", self.billboard),
                ("Clearing/cleaning charges", self.cleaning),
                ("Miscellaneous expenses", self.misc_expenses),
                ("Federation fee", self.federation_fee),
                ("HOA-BOD Uniforms", self.uniforms),
                ("BOD Mtg", self.bod_mtg),
                ("General Assembly", self.general_assembly),
                ("Cash Deposit to bank", self.cash_deposit),
                ("Withholding tax on bank deposit", self.withholding_tax),
                ("Refund for seri-culture", self.refund_sericulture),
                ("Others", self.outflows_others),
                ("", self.outflows_others_2)
            ]
            for i, (label, var) in enumerate(outflow_items):
                table.cell(i, 0).text = label
                table.cell(i, 1).text = format_amount(var.get())
                table.cell(i, 1).paragraphs[0].alignment = 2  # Right align amounts
            table.cell(0, 0).paragraphs[0].runs[0].bold = True  # Bold total

            # Ending Cash Balance
            doc.add_heading("Ending Cash Balance", level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.cell(0, 0).text = "Ending cash balance"
            table.cell(0, 1).text = format_amount(self.ending_cash.get())
            table.cell(0, 1).paragraphs[0].alignment = 2  # Right align
            table.cell(0, 0).paragraphs[0].runs[0].bold = True  # Bold

            # Breakdown of Cash
            doc.add_heading("Breakdown of Cash", level=2)
            table = doc.add_table(rows=2, cols=2)
            table.style = 'Table Grid'
            table.cell(0, 0).text = "Cash in Bank"
            table.cell(0, 1).text = format_amount(self.ending_cash_bank.get())
            table.cell(1, 0).text = "Cash on Hand"
            table.cell(1, 1).text = format_amount(self.ending_cash_hand.get())
            for row in table.rows:
                row.cells[1].paragraphs[0].alignment = 2  # Right align amounts

            # Save the document
            doc.save(filename)
            messagebox.showinfo("Success", f"Word document saved to {filename}")
            return filename  # Return filename for email attachment
            
        except Exception as e:
            messagebox.showerror("Error", f"Error saving to Word: {str(e)}\n\nMake sure you have python-docx installed by running:\npip install python-docx")
            return None

    def send_email(self):
        try:
            # Get email details from GUI
            sender_email = self.sender_email_var.get().strip()
            sender_password = self.sender_password_var.get().strip()
            recipient_emails = [email.strip() for email in self.recipient_emails_var.get().split(',') if email.strip()]
            
            if not sender_email or not sender_password or not recipient_emails:
                messagebox.showerror("Error", "Please fill in all email fields (Sender Email, App Password, Recipients).")
                return

            # Create PDF and Word files to attach (both auto-calculate totals)
            pdf_filename = self.export_to_pdf()
            if not pdf_filename:
                return
            docx_filename = self.save_to_docx()
            if not docx_filename:
                os.remove(pdf_filename)  # Clean up PDF if Word fails
                return

            # Set up the email
            msg = MIMEMultipart()
            msg['From'] = sender_email
            msg['To'] = ", ".join(recipient_emails)
            msg['Subject'] = f"Cash Flow Statement - {self.today_date}"

            body = f"Attached is the cash flow statement for {self.today_date} in both PDF and Word formats.\n\nRegards,\nYour Name"
            msg.attach(MIMEText(body, 'plain'))

            # Attach PDF
            with open(pdf_filename, 'rb') as f:
                part = MIMEApplication(f.read(), Name=os.path.basename(pdf_filename))
                part['Content-Disposition'] = f'attachment; filename="{os.path.basename(pdf_filename)}"'
                msg.attach(part)

            # Attach Word
            with open(docx_filename, 'rb') as f:
                part = MIMEApplication(f.read(), Name=os.path.basename(docx_filename))
                part['Content-Disposition'] = f'attachment; filename="{os.path.basename(docx_filename)}"'
                msg.attach(part)

            # Connect to SMTP server (assuming Gmail)
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login(sender_email, sender_password)
            
            server.send_message(msg)
            server.quit()
            
            messagebox.showinfo("Success", f"Email with PDF and Word files sent to {', '.join(recipient_emails)}!")
            os.remove(pdf_filename)  # Clean up temporary PDF
            os.remove(docx_filename)  # Clean up temporary Word file
            
        except smtplib.SMTPAuthenticationError as e:
            messagebox.showerror("Error", f"Authentication failed: {str(e)}\nCheck your email and app password.")
            if 'pdf_filename' in locals():
                os.remove(pdf_filename)
            if 'docx_filename' in locals():
                os.remove(docx_filename)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to send email: {str(e)}\nEnsure your email credentials are correct and you have an internet connection.")
            if 'pdf_filename' in locals():
                os.remove(pdf_filename)
            if 'docx_filename' in locals():
                os.remove(docx_filename)

if __name__ == "__main__":
    root = tk.Tk()
    app = IntegratedCashFlowApp(root)
    root.mainloop()